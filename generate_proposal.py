from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

def set_font(run, name='宋体', size=10.5, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color

def add_title(text, level=0):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        set_font(run, '黑体', 22 if level==0 else 16, True, RGBColor(0x2C,0x3E,0x50))
    return p

def add_section(title_text):
    h = doc.add_heading(title_text, level=1)
    for run in h.runs:
        set_font(run, '黑体', 16, True, RGBColor(0x2C,0x3E,0x50))

def add_para(text, bold=False):
    p = doc.add_paragraph(text)
    for run in p.runs:
        set_font(run, '宋体', 11, bold)
    return p

# 封面
add_title('基于人工智能的老年衰弱')
add_title('早期识别与干预系统研发', level=1)
add_title('院级科研课题申请书', level=1)
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run('申请人：__________'), '宋体', 12)
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run('申请科室：__________'), '宋体', 12)
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run('申请日期：2026年____月____日'), '宋体', 12)

doc.add_page_break()

# 一、立项依据
add_section('一、立项依据与研究意义')

add_section('1.1 研究背景')
add_para('衰弱（Frailty）是一种与年龄相关的生理储备下降状态，表现为多系统功能减退、应激耐受力降低。老年衰弱患者面临更高的跌倒、住院、失能和死亡风险。据文献报道，社区老年人衰弱患病率约为10%-15%，住院老年患者可达30%以上。')
add_para('目前临床对衰弱的识别主要依赖医生的经验判断，缺乏标准化、量化的评估工具。Fried衰弱表型标准虽为国际公认，但需测量体重变化、握力、步速、活动量等5项指标，临床操作繁琐，且不同医生判断标准不一，导致漏诊率和误诊率较高。')
add_para('随着人工智能技术的发展，自然语言处理（NLP）和大型语言模型（LLM）在医学文本分析领域展现出巨大潜力。利用电子病历文本，结合深度学习模型自动识别衰弱状态，有望实现早期、快速、标准化的衰弱筛查。')

add_section('1.2 国内外研究现状')
add_para('国外研究：')
add_para('• 2022年，Johns Hopkins大学研究团队利用电子健康记录（EHR）数据，结合机器学习模型预测老年衰弱，AUC达0.85')
add_para('• 2023年，Nature Aging发表研究，利用Transformer模型分析临床笔记，实现衰弱自动识别')
add_para('• 欧美多国已将衰弱评估纳入老年综合评估（CGA）标准流程')
add_para('')
add_para('国内研究：')
add_para('• 国内衰弱研究起步较晚，主要集中在流行病学调查和干预研究')
add_para('• 人工智能在衰弱识别领域的应用研究较少，尚无成熟的临床落地产品')
add_para('• 现有研究多依赖结构化数据，对非结构化病历文本的利用不足')
add_para('')
add_para('现有不足：')
add_para('1. 缺乏针对中文病历的衰弱识别专用模型')
add_para('2. 现有模型多为"黑箱"，缺乏可解释性，临床接受度低')
add_para('3. 模型训练数据多来自国外，与我国老年人群特征存在差异')
add_para('4. 缺乏持续学习机制，模型难以适应本院临床实际')

add_section('1.3 研究意义')
add_para('理论意义：')
add_para('• 探索基于中文电子病历的衰弱自动识别方法，填补国内研究空白')
add_para('• 构建"BERT分类 + RAG检索增强 + LLM解释"的多模态融合架构，为医学AI可解释性研究提供新思路')
add_para('• 建立医生反馈驱动的持续学习机制，为医学AI模型的临床迭代优化提供方法论参考')
add_para('')
add_para('临床意义：')
add_para('• 提高衰弱识别效率：从人工逐项评估（约10-15分钟/例）到系统自动分析（约5秒/例）')
add_para('• 降低漏诊率：系统可24小时不间断筛查，不放过任何可疑病例')
add_para('• 促进标准化：统一评估标准，减少不同医生间的主观差异')
add_para('• 辅助临床决策：提供结构化评估报告和循证建议，提升诊疗质量')
add_para('• 支持科研：积累标准化衰弱数据，为后续临床研究提供数据基础')

add_section('1.4 参考文献')
add_para('[1] Fried LP, et al. Frailty in older adults: evidence for a phenotype. J Gerontol A Biol Sci Med Sci. 2001;56(3):M146-M156.')
add_para('[2] Rockwood K, et al. A global clinical measure of fitness and frailty in elderly people. CMAJ. 2005;173(5):489-495.')
add_para('[3] 于普林, 等. 中国老年患者衰弱评估与干预专家共识. 中华老年医学杂志. 2023;42(1):1-10.')
add_para('[4] Devlin J, et al. BERT: Pre-training of Deep Bidirectional Transformers for Language Understanding. NAACL 2019.')
add_para('[5] Lewis P, et al. Retrieval-Augmented Generation for Knowledge-Intensive NLP Tasks. NeurIPS 2020.')
add_para('[6] 中华医学会老年医学分会. 老年衰弱门诊建设专家共识. 中华老年医学杂志. 2024.')

doc.add_page_break()

# 二、研究目标与内容
add_section('二、研究目标与内容')

add_section('2.1 总体目标')
add_para('研发一套基于人工智能的老年衰弱早期识别与干预系统，实现从电子病历文本到衰弱评估结论的自动化、标准化、可解释化，并建立医生反馈驱动的持续学习机制，使模型性能随临床使用不断优化。')

add_section('2.2 具体目标')
add_para('目标1：构建基于中文电子病历的衰弱识别BERT模型，准确率≥85%')
add_para('目标2：建立本地医学知识库，实现RAG检索增强的LLM解释生成')
add_para('目标3：开发医生反馈收集机制，建立持续学习闭环')
add_para('目标4：完成本院临床验证，评估系统的临床可用性和接受度')
add_para('目标5：形成标准化衰弱数据集，发表相关论文1-2篇')

add_section('2.3 研究内容')
add_para('内容1：数据采集与预处理')
add_para('• 收集本院老年科近3年脱密电子病历500-1000例')
add_para('• 由2名高年资医生独立标注衰弱状态，不一致时由第3人仲裁')
add_para('• 数据清洗：去除敏感信息、标准化术语、补齐缺失字段')
add_para('')
add_para('内容2：模型研发')
add_para('• 基于bert-base-chinese构建二分类模型（衰弱/非衰弱）')
add_para('• 构建ChromaDB向量知识库，收录老年医学指南和文献')
add_para('• 集成本地qwen2.5:3b大模型，实现结构化评估报告生成')
add_para('• 开发医生反馈对话框，收集校正意见和原因')
add_para('')
add_para('内容3：持续学习机制')
add_para('• 阶段1（0-20条分歧）：数据积累期')
add_para('• 阶段2（≥20条）：自动更新LLM Prompt，融入历史错误案例')
add_para('• 阶段3（≥50条）：后台增量微调BERT，生成新版本模型')
add_para('• 阶段4（≥100条）：启用动态相似案例检索，个性化提示')
add_para('')
add_para('内容4：临床验证')
add_para('• 前瞻性收集100例老年患者病历，对比系统评估与专家评估的一致性')
add_para('• 计算敏感度、特异度、Kappa一致性系数')
add_para('• 调查医生对系统的满意度（Likert 5级量表）')
add_para('• 评估系统对临床工作流程的影响')

doc.add_page_break()

# 三、技术路线
add_section('三、技术路线与实施方案')

add_section('3.1 系统架构')
add_para('本系统采用"端-边-云"协同架构：')
add_para('• 端：PySide6桌面应用，医生操作界面')
add_para('• 边：本地Ollama服务，运行qwen2.5:3b大模型')
add_para('• 数据：本地SQLite数据库 + JSONL日志，确保数据安全')
add_para('')
add_para('核心技术栈：')
add_para('• 分类模型：BERT (bert-base-chinese)')
add_para('• 向量检索：ChromaDB + BGE-large-zh-v1.5')
add_para('• 解释生成：RAG + Ollama (qwen2.5:3b)')
add_para('• GUI框架：PySide6')
add_para('• 数据存储：SQLite + JSONL（双保险）')

add_section('3.2 技术路线')
add_para('第一阶段：基础研发（1-3个月）')
add_para('• 完成BERT模型训练和基础评估功能')
add_para('• 构建ChromaDB知识库，收录老年医学指南')
add_para('• 开发PySide6桌面应用，实现单病例评估')
add_para('')
add_para('第二阶段：功能完善（4-6个月）')
add_para('• 开发批量导入和批量标注功能')
add_para('• 实现医生反馈收集和持续学习机制')
add_para('• 完成病历缺失指标警告和LLM重试机制')
add_para('')
add_para('第三阶段：临床验证（7-12个月）')
add_para('• 收集本院病历数据，进行模型训练和调优')
add_para('• 开展前瞻性临床验证研究')
add_para('• 收集医生反馈，优化系统界面和功能')
add_para('• 撰写论文，申请软件著作权')

add_section('3.3 创新点')
add_para('创新点1：中文病历衰弱识别专用模型')
add_para('• 针对中文电子病历特点，基于bert-base-chinese构建衰弱识别模型')
add_para('• 融合Fried标准5项指标的关键特征，提升识别准确率')
add_para('')
add_para('创新点2：可解释性AI架构')
add_para('• 采用"BERT分类 + RAG检索 + LLM解释"三层架构')
add_para('• 不仅给出结论，还提供逐项核对和循证建议')
add_para('• 医生可理解、可验证、可信任')
add_para('')
add_para('创新点3：医生反馈驱动的持续学习')
add_para('• 建立"使用-反馈-优化"闭环，模型性能随临床使用不断提升')
add_para('• 自动触发Prompt更新、增量微调、相似案例检索三阶段优化')
add_para('• 模型热切换，无需重启，不影响临床工作')
add_para('')
add_para('创新点4：数据安全与隐私保护')
add_para('• 所有数据本地存储，不上传云端')
add_para('• 病历脱敏处理，去除姓名、身份证号等敏感信息')
add_para('• JSONL日志实时刷盘，系统崩溃不丢失数据')

doc.add_page_break()

# 四、预期成果
add_section('四、预期成果与考核指标')

add_section('4.1 预期成果')
add_para('1. 软件系统：老年衰弱智能评估系统v1.0（含源代码）')
add_para('2. 数据集：标准化老年衰弱标注数据集（≥500例）')
add_para('3. 论文：中文核心期刊或SCI论文1-2篇')
add_para('4. 知识产权：软件著作权1项')
add_para('5. 临床报告：系统临床验证研究报告')

add_section('4.2 考核指标')
add_para('技术指标：')
add_para('• BERT模型准确率≥85%，敏感度≥80%，特异度≥85%')
add_para('• LLM输出格式正确率≥90%')
add_para('• 单病例评估时间≤10秒')
add_para('• 系统运行稳定性≥99%（无崩溃）')
add_para('')
add_para('临床指标：')
add_para('• 与专家评估一致性Kappa≥0.75')
add_para('• 医生满意度评分≥4.0/5.0')
add_para('• 衰弱识别漏诊率降低≥20%')
add_para('')
add_para('应用指标：')
add_para('• 在本院老年科推广应用')
add_para('• 累计使用病例≥1000例')
add_para('• 收集医生反馈≥200条')

doc.add_page_break()

# 五、研究基础
add_section('五、研究基础与条件')

add_section('5.1 研究基础')
add_para('申请人长期从事老年医学临床工作，对老年衰弱评估有丰富经验。')
add_para('课题组已初步完成系统原型开发，具备以下基础：')
add_para('• BERT分类模型：基于bert-base-chinese，已完成初步训练')
add_para('• 知识库：已构建ChromaDB向量库，收录老年医学指南')
add_para('• 桌面应用：PySide6 GUI已完成，支持单病例评估和批量导入')
add_para('• 反馈机制：已实现医生反馈收集和JSONL防丢失日志')
add_para('• 持续学习：已开发自动调度器，支持四阶段优化')

add_section('5.2 研究条件')
add_para('硬件条件：')
add_para('• 工作站：配备NVIDIA GPU（显存≥8GB），满足模型训练和推理需求')
add_para('• 服务器：用于Ollama本地大模型服务')
add_para('• 存储：≥500GB SSD，满足模型和数据存储')
add_para('')
add_para('软件条件：')
add_para('• Python 3.12 + PyTorch + Transformers')
add_para('• Ollama本地大模型运行环境')
add_para('• PySide6桌面应用框架')
add_para('')
add_para('数据条件：')
add_para('• 本院电子病历系统已运行多年，数据积累丰富')
add_para('• 已与信息科沟通，可导出脱敏病历数据')
add_para('• 老年科年门诊量约5000人次，住院量约2000人次，病例来源充足')
add_para('')
add_para('团队条件：')
add_para('• 申请人：老年医学副主任医师，主攻老年衰弱和综合评估')
add_para('• 技术顾问：软件公司工程师，负责AI模型开发和系统维护')
add_para('• 协作医生：2名高年资主治医师，负责病历标注和质量审核')
add_para('• 统计学支持：医院统计室，负责数据分析和论文撰写')

doc.add_page_break()

# 六、经费预算
add_section('六、经费预算')

add_para('总预算：____万元')
add_para('')
add_para('1. 设备费：____万元')
add_para('   • GPU工作站升级：____万元')
add_para('   • 数据存储设备：____万元')
add_para('')
add_para('2. 材料费：____万元')
add_para('   • 病历数据脱敏处理：____万元')
add_para('   • 标注劳务费（500例×50元）：____万元')
add_para('')
add_para('3. 测试化验加工费：____万元')
add_para('   • 软件测试和第三方评估：____万元')
add_para('')
add_para('4. 差旅费：____万元')
add_para('   • 学术会议交流（2人次）：____万元')
add_para('')
add_para('5. 会议费：____万元')
add_para('   • 项目启动会和结题会：____万元')
add_para('')
add_para('6. 出版/文献费：____万元')
add_para('   • 论文发表版面费：____万元')
add_para('   • 软件著作权申请：____万元')
add_para('')
add_para('7. 劳务费：____万元')
add_para('   • 研究生/实习生劳务：____万元')
add_para('')
add_para('8. 其他：____万元')
add_para('   • 不可预见费：____万元')

doc.add_page_break()

# 七、进度安排
add_section('七、进度安排')

add_para('第一年（1-6月）：基础研发阶段')
add_para('• 完成病历数据收集和预处理（500例）')
add_para('• 完成BERT模型训练和调优')
add_para('• 完成知识库构建和RAG检索功能')
add_para('• 完成桌面应用基础功能开发')
add_para('')
add_para('第一年（7-12月）：功能完善阶段')
add_para('• 开发批量标注和模型重训练功能')
add_para('• 实现医生反馈收集和持续学习机制')
add_para('• 完成系统内部测试和bug修复')
add_para('• 申请软件著作权')
add_para('')
add_para('第二年（1-6月）：临床验证阶段')
add_para('• 开展前瞻性临床验证研究（100例）')
add_para('• 收集医生反馈，优化系统')
add_para('• 撰写临床验证报告')
add_para('• 撰写论文初稿')
add_para('')
add_para('第二年（7-12月）：总结推广阶段')
add_para('• 完成论文投稿和修改')
add_para('• 在本院老年科推广应用')
add_para('• 撰写结题报告')
add_para('• 申请后续省级/国家级课题')

doc.add_page_break()

# 八、风险分析
add_section('八、风险分析与应对措施')

add_para('风险1：模型准确率不达标')
add_para('• 应对措施：增加训练数据量、尝试更大模型（如bert-large-chinese）、引入多模型集成')
add_para('')
add_para('风险2：病历数据质量不佳')
add_para('• 应对措施：与信息科协作规范病历书写、建立数据清洗流程、设置缺失指标警告')
add_para('')
add_para('风险3：医生接受度低')
add_para('• 应对措施：优化界面设计、提供详细解释、建立反馈机制、开展培训')
add_para('')
add_para('风险4：数据安全和隐私泄露')
add_para('• 应对措施：严格脱敏处理、本地存储不上云、签署数据使用协议、定期安全审计')
add_para('')
add_para('风险5：Ollama服务不稳定')
add_para('• 应对措施：备用模型方案、本地缓存机制、降级到BERT单独评估')

doc.add_page_break()

# 九、伦理与数据安全
add_section('九、伦理审查与数据安全')

add_para('9.1 伦理审查')
add_para('• 本研究仅使用已完成诊疗的病历数据，不涉及额外人体试验')
add_para('• 所有病历数据已完成脱敏处理，去除姓名、身份证号、联系方式等可识别信息')
add_para('• 将向医院伦理委员会提交研究方案，获取伦理审批')
add_para('• 研究过程遵循《赫尔辛基宣言》和《涉及人的生物医学研究伦理审查办法》')
add_para('')
add_para('9.2 数据安全')
add_para('• 数据存储：所有数据保存在医院内网服务器，不上传互联网')
add_para('• 访问控制：设置分级权限，仅课题组成员可访问')
add_para('• 数据备份：定期备份数据库，防止数据丢失')
add_para('• 使用范围：仅用于本课题研究，不用于商业用途')
add_para('• 数据销毁：研究结束后，按医院规定销毁原始数据')

doc.add_page_break()

# 十、附件清单
add_section('十、附件清单')

add_para('□ 申请人简历及代表性成果')
add_para('□ 课题组主要成员简历')
add_para('□ 前期研究成果证明材料')
add_para('□ 系统原型截图或演示视频')
add_para('□ 信息科数据使用支持函')
add_para('□ 伦理审查申请表')
add_para('□ 经费预算明细表')
add_para('□ 合作单位协议（如有）')

# 保存
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '院级课题申请书_老年衰弱AI识别系统.docx')
doc.save(output_path)
print(f'已保存: {output_path}')
print(f'文件大小: {os.path.getsize(output_path)} bytes')
