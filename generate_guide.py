from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

def set_font(run, name='宋体', size=10.5, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color

def add_title(text, level=0):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        set_font(run, '黑体', 22 if level==0 else 16, True, RGBColor(0x2C,0x3E,0x50))
    return p

def add_section(title_text):
    h = doc.add_heading(title_text, level=1)
    for run in h.runs:
        set_font(run, '黑体', 16, True, RGBColor(0x2C,0x3E,0x50))

def add_para(text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        set_font(run, '宋体', 11)
    return p

# 封面
add_title('老年衰弱智能评估系统')
add_title('用户指南', level=1)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run('基于 BERT + RAG + 本地大模型的老年衰弱自动识别系统'), '楷体', 12, False, RGBColor(0x7F,0x8C,0x8D))
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run('版本：v2.0  |  更新日期：2026-07-20'), '宋体', 10, False, RGBColor(0x95,0xA5,0xA6))

doc.add_page_break()

# 目录
add_section('目录')
items = [
    '一、系统概述', '二、环境准备与启动', '三、主界面功能介绍',
    '四、单病例评估流程', '五、批量导入病例', '六、批量标注与模型训练',
    '七、反馈收集与持续学习', '八、数据导出', '九、常见问题', '十、注意事项'
]
for item in items:
    p = doc.add_paragraph(item, style='List Number')
    for run in p.runs:
        set_font(run, '宋体', 11)

doc.add_page_break()

# 一、系统概述
add_section('一、系统概述')
add_para('本系统基于电子病历文本，自动识别老年患者是否存在衰弱（Frailty）状态，并提供结构化临床解释建议。')
add_para('核心技术：')
add_para('• BERT 分类模型：基于 bert-base-chinese，二分类（衰弱 / 非衰弱）')
add_para('• RAG 检索增强：从本地知识库检索相关指南，辅助 LLM 生成解释')
add_para('• 本地大模型：通过 Ollama 运行 qwen2.5:3b，无需联网')
add_para('• 持续学习：收集医生反馈，自动优化模型和 Prompt')
add_para('')
add_para('评估标准：采用 Fried 衰弱表型标准（5项，满足≥3项为衰弱）：')
add_para('1. 不明原因体重下降（过去1年内体重下降≥5%或≥4.5kg）')
add_para('2. 疲乏/精力减退（自我报告"疲乏无力"、"整天没精神"等）')
add_para('3. 握力下降（用手握力计测量，根据性别和BMI调整阈值）')
add_para('4. 步速减慢（行走4.57米时间延长，根据身高、性别调整）')
add_para('5. 活动量减少（每周消耗热量低于标准）')

# 二、环境准备与启动
add_section('二、环境准备与启动')
add_para('1. 启动 Ollama 服务：在命令行执行 ollama serve')
add_para('2. 确保模型已下载：ollama pull qwen2.5:3b（首次约2GB）')
add_para('3. 启动主程序：')
add_para('   cd pythonProject/source')
add_para('   python frailty_assessment_app.py')
add_para('4. 首次启动会自动加载 BERT 模型（约400MB）和 BGE 嵌入模型（约1.3GB）')

# 三、主界面功能介绍
add_section('三、主界面功能介绍')
add_para('【文件(F)】菜单：')
add_para('• 加载病例... — 打开单个 txt 病历文件')
add_para('• 批量导入... — 批量导入 CSV/Excel/txt 并自动评估')
add_para('• 导出 Excel — 将患者列表导出为表格')
add_para('• 退出 — 关闭程序')
add_para('')
add_para('【工具(T)】菜单：')
add_para('• 批量标注与模型训练... — 打开批量标注与模型重训练面板')
add_para('• 模型信息 / Ollama 状态 / 系统设置')
add_para('')
add_para('【帮助(H)】菜单：')
add_para('• 关于软件 / 使用指南')
add_para('')
add_para('工具栏：快速预测、详细解释、清空、加载病例')
add_para('')
add_para('结果区：')
add_para('• 快速预测标签页：BERT 二分类结果和置信度')
add_para('• 详细解释标签页：LLM 生成的结构化报告（Fried 5项逐项核对、阳性计数、结论、建议）')

# 四、单病例评估流程
add_section('四、单病例评估流程')
add_para('步骤1：在左侧文本框粘贴完整病历文本（主诉、现病史、既往史、查体结果）')
add_para('步骤2：点击【快速预测】，查看 BERT 分类结果')
add_para('步骤3：点击【详细解释】，查看 LLM 生成的结构化评估报告')
add_para('步骤4（如结果有误）：点击【💡 提交反馈】→ 选择正确结论 → 填写原因 → 保存')
add_para('')
add_para('病历缺失指标警告：')
add_para('如果病历缺少≥3项 Fried 关键指标，系统会弹出警告。可选择"是"继续（结果可能不可靠）或"否"补充病历后再试。')

# 五、批量导入病例
add_section('五、批量导入病例')
add_para('1. 点击【文件】→【批量导入...】')
add_para('2. 选择病历文件（可多选），支持 .txt / .csv / .xlsx / .xls')
add_para('3. 系统自动识别列名：病历文本、text、主诉现病史、入院记录、病情描述')
add_para('4. 导入完成后，患者列表表格显示所有评估结果')
add_para('5. 双击表格行可加载该病历到输入区')
add_para('6. 姓名列根据预测结果着色：红色=衰弱、橙色=衰弱前期、绿色=非衰弱')

# 六、批量标注与模型训练
add_section('六、批量标注与模型训练')
add_para('适用场景：系统上线前冷启动 / 定期模型大补')
add_para('')
add_para('步骤1：导入脱密病历 → 点击【选择 CSV/Excel 文件】→ 自动 BERT 预标注')
add_para('步骤2：LLM 预标注 → 点击【运行 LLM 预标注】→ 对不确定样本详细评估')
add_para('步骤3：人工审核 → 点击【加载待审核记录】→ 逐条审核（黄色高亮=不确定样本）')
add_para('         点击表格行，右侧显示完整病历原文 → 选择【衰弱/衰弱前期/非衰弱】')
add_para('步骤4：数据分析 → 点击【查看标注质量分析】')
add_para('步骤5：模型重训练 → 设置最少样本和 Epoch → 点击【全量重训练 BERT】')
add_para('         训练完成后可选择是否切换到新模型')
add_para('步骤6：数据导出 → 点击【导出标注数据（CSV）】')
add_para('')
add_para('与反馈收集的区别：')
add_para('• 批量标注：历史脱密病历 → 批量预标注+人工审核 → 全量重训练')
add_para('• 反馈收集：日常使用中发现的错误 → 逐条提交 → 增量微调')

# 七、反馈收集与持续学习
add_section('七、反馈收集与持续学习')
add_para('反馈收集：完成详细解释后，点击【💡 提交反馈】→ 选择正确结论 → 保存')
add_para('')
add_para('数据安全：双保险存储（SQLite + JSONL 日志），每写一条立即强制刷盘，系统崩溃不丢失。')
add_para('')
add_para('自动持续学习（根据反馈分歧数自动触发）：')
add_para('• 阶段1（0-20条）：只收集')
add_para('• 阶段2（≥20条）：自动更新 LLM Prompt')
add_para('• 阶段3（≥50条）：后台自动增量微调 BERT')
add_para('• 阶段4（≥100条）：自动启用动态相似案例嵌入')
add_para('')
add_para('模型热切换：新模型训练完成后无需重启，点击确认即可切换，原模型保留可回退。')

# 八、数据导出
add_section('八、数据导出')
add_para('1. 导出患者列表：【文件】→【导出 Excel】')
add_para('2. 导出反馈数据：Python 中执行 from frailty_assessment_app import export_feedback_to_csv; export_feedback_to_csv("path.csv")')
add_para('3. 导出批量标注数据：在批量标注对话框中点击【导出标注数据（CSV）】')

# 九、常见问题
add_section('九、常见问题')
add_para('Q1：启动提示"无法连接到 Ollama"？')
add_para('A：确保已运行 ollama serve，且 qwen2.5:3b 已下载。')
add_para('')
add_para('Q2：详细解释生成失败？')
add_para('A：3B 模型能力有限，系统会自动重试 3 次。检查 Ollama 服务状态。')
add_para('')
add_para('Q3：批量导入提示"未找到病历文本列"？')
add_para('A：确保 CSV/Excel 包含列名：病历文本、text、主诉现病史、入院记录、病情描述（之一即可）')
add_para('')
add_para('Q4：训练提示"样本不足"？')
add_para('A：全量重训练默认需要至少 100 条已标注数据。请先人工审核足够数量。')
add_para('')
add_para('Q5：如何清理测试数据？')
add_para('A：删除 batch_annotation.db、feedback.db、feedback_log.jsonl、frailty_bert_retrained_v*/、frailty_bert_retrain/ 即可恢复初始状态。')
add_para('')
add_para('Q6：系统支持哪些病历格式？')
add_para('A：单个 .txt；批量 .csv / .xlsx / .xls（需含病历文本列）')

# 十、注意事项
add_section('十、注意事项')
add_para('1. 数据隐私：所有数据本地存储，不上传云端。导入病历应已完成脱密处理。')
add_para('2. 模型局限：BERT 仅在演示数据训练，真实场景准确率需验证；LLM（3B）能力有限，输出仅供参考，最终诊断由医生确认。')
add_para('3. 使用建议：初期以收集反馈为主，先用小批量真实病历验证流程，确认无误后再大规模导入。')
add_para('4. 定期备份 feedback.db 和 batch_annotation.db。')
add_para('5. 如有问题，查看项目文档.md 或 data/问题记录.md。')

# 保存
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '老年衰弱智能评估系统_用户指南.docx')
doc.save(output_path)
print(f'已保存: {output_path}')
print(f'文件大小: {os.path.getsize(output_path)} bytes')
