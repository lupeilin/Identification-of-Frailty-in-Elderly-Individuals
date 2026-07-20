from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

def set_font(run, name='宋体', size=10.5, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color

def add_title(text, level=0):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        set_font(run, '黑体', 22 if level==0 else 16, True, RGBColor(0x2C,0x3E,0x50))
    return p

def add_section(title_text):
    h = doc.add_heading(title_text, level=1)
    for run in h.runs:
        set_font(run, '黑体', 16, True, RGBColor(0x2C,0x3E,0x50))

def add_subsection(title_text):
    h = doc.add_heading(title_text, level=2)
    for run in h.runs:
        set_font(run, '黑体', 13, True, RGBColor(0x2C,0x3E,0x50))

def add_para(text, bold=False):
    p = doc.add_paragraph(text)
    for run in p.runs:
        set_font(run, '宋体', 11, bold)
    return p

# 封面
add_title('基于人工智能的老年衰弱识别系统')
add_title('临床验证与效果评价研究方案', level=1)
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run('—— 中文核心期刊投稿用实验设计方案 ——'), '楷体', 12, False, RGBColor(0x7F,0x8C,0x8D))
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run('目标期刊：《中华老年医学杂志》或《中国老年学杂志》'), '宋体', 11, False, RGBColor(0x95,0xA5,0xA6))

doc.add_page_break()

# 目录
add_section('目录')
items = [
    '一、研究背景与目的',
    '二、研究设计',
    '三、研究对象',
    '四、研究方法',
    '五、评价指标',
    '六、样本量计算',
    '七、统计分析',
    '八、质量控制',
    '九、伦理考虑',
    '十、论文框架',
    '十一、时间进度',
    '十二、预期成果',
]
for item in items:
    p = doc.add_paragraph(item, style='List Number')
    for run in p.runs:
        set_font(run, '宋体', 11)

doc.add_page_break()

# 一、研究背景与目的
add_section('一、研究背景与目的')

add_subsection('1.1 研究背景')
add_para('衰弱（Frailty）是一种与年龄相关的生理储备下降状态，表现为多系统功能减退、应激耐受力降低。老年衰弱患者面临更高的跌倒、住院、失能和死亡风险。据文献报道，社区老年人衰弱患病率约为10%-15%，住院老年患者可达30%以上。')
add_para('目前临床对衰弱的识别主要依赖医生的经验判断，缺乏标准化、量化的评估工具。Fried衰弱表型标准虽为国际公认，但需测量体重变化、握力、步速、活动量等5项指标，临床操作繁琐，且不同医生判断标准不一，导致漏诊率和误诊率较高。')
add_para('近年来，人工智能技术在医学文本分析领域展现出巨大潜力。基于自然语言处理（NLP）的深度学习模型可以从电子病历文本中自动提取临床特征，辅助疾病识别。然而，目前针对中文电子病历的老年衰弱自动识别研究较少，且现有模型多为"黑箱"，缺乏可解释性，临床接受度低。')
add_para('本研究团队在既往工作基础上，开发了一套基于BERT+RAG+本地大模型的老年衰弱智能评估系统，实现了从电子病历文本到衰弱评估结论的自动化、标准化、可解释化。本研究旨在通过临床验证，评估该系统在真实临床环境中的识别准确性和实用性。')

add_subsection('1.2 研究目的')
add_para('1. 评估基于人工智能的老年衰弱识别系统对中文电子病历的衰弱识别准确性。')
add_para('2. 比较人工智能系统与临床医生诊断的一致性。')
add_para('3. 比较人工智能系统与FRAIL量表的筛查效果。')
add_para('4. 评估医生对该系统的使用满意度和接受度。')

doc.add_page_break()

# 二、研究设计
add_section('二、研究设计')

add_subsection('2.1 设计类型')
add_para('诊断准确性研究（Diagnostic Accuracy Study），采用横断面设计。')

add_subsection('2.2 研究思路')
add_para('1. 收集老年患者的电子病历文本和临床资料。')
add_para('2. 由2名高年资老年科医生独立进行衰弱评估（Fried表型标准），作为"金标准"。')
add_para('3. 由人工智能系统对同一批病历进行自动评估。')
add_para('4. 由第3名医生使用FRAIL量表进行快速筛查。')
add_para('5. 比较人工智能系统、FRAIL量表与"金标准"的一致性。')
add_para('6. 对使用系统的医生进行满意度调查。')

doc.add_page_break()

# 三、研究对象
add_section('三、研究对象')

add_subsection('3.1 纳入标准')
add_para('1. 年龄≥65岁。')
add_para('2. 在某医院老年科住院或门诊就诊的患者。')
add_para('3. 电子病历资料完整，包含主诉、现病史、既往史、查体记录。')
add_para('4. 患者或家属知情同意，自愿参与本研究。')

add_subsection('3.2 排除标准')
add_para('1. 病历资料严重缺失，无法进行衰弱评估。')
add_para('2. 急性疾病期（如急性心肌梗死、急性脑卒中、急性感染等），病情不稳定。')
add_para('3. 终末期疾病，预期寿命<3个月。')
add_para('4. 患者或家属拒绝参与。')

add_subsection('3.3 样本量')
add_para('计划纳入200例老年患者。')
add_para('（详见第六部分样本量计算）')

doc.add_page_break()

# 四、研究方法
add_section('四、研究方法')

add_subsection('4.1 数据收集')
add_para('1. 从医院电子病历系统导出符合纳入标准的患者病历。')
add_para('2. 病历数据脱敏处理：去除患者姓名、身份证号、手机号等可识别信息，用编码代替。')
add_para('3. 收集的临床资料包括：')
add_para('   - 基本信息：年龄、性别、身高、体重、BMI。')
add_para('   - 病历文本：主诉、现病史、既往史、个人史、家族史、查体记录、辅助检查。')
add_para('   - 实验室检查：血常规、肝肾功能、电解质、白蛋白、前白蛋白、维生素D、甲状腺功能等。')
add_para('   - 功能评估：日常生活能力（ADL）、工具性日常生活能力（IADL）。')

add_subsection('4.2 金标准评估（医生诊断）')
add_para('由2名具有5年以上老年科临床经验的副主任医师独立进行衰弱评估。')
add_para('评估工具：Fried衰弱表型标准（5项指标）。')
add_para('评估方法：')
add_para('  - 仔细阅读病历文本，提取与5项指标相关的信息。')
add_para('  - 必要时补充测量：握力（Jamar握力计）、步速（4.57米步行测试）。')
add_para('  - 根据5项指标判定：0项=非衰弱，1-2项=衰弱前期，≥3项=衰弱。')
add_para('2名医生评估结果不一致时，由第3名主任医师仲裁，以仲裁结果作为最终金标准。')

add_subsection('4.3 人工智能系统评估')
add_para('将脱敏后的病历文本输入人工智能系统，系统自动输出：')
add_para('  - BERT预测结果（衰弱/非衰弱）和置信度。')
add_para('  - LLM生成的结构化评估报告（Fried 5项逐项核对、阳性计数、结论、临床建议）。')
add_para('系统评估过程完全自动化，评估人员不参与，避免人为干预。')

add_subsection('4.4 FRAIL量表筛查')
add_para('由第3名医生（不参与金标准评估）使用FRAIL量表进行快速筛查。')
add_para('FRAIL量表包含5个问题，每个"是"得1分：')
add_para('  1. 您是否感到疲劳？')
add_para('  2. 您是否不能爬上1层楼梯？')
add_para('  3. 您是否不能走100米？')
add_para('  4. 您是否患有5种以上疾病？')
add_para('  5. 您是否最近1年内体重下降超过5%？')
add_para('评分：0分=健康，1-2分=衰弱前期，3-5分=衰弱。')

add_subsection('4.5 人工智能系统介绍')
add_para('本系统采用"BERT分类 + RAG检索增强 + 本地大模型解释"三层架构：')
add_para('1. BERT分类模型：基于bert-base-chinese，对病历文本进行二分类（衰弱/非衰弱）。')
add_para('2. RAG检索增强：从本地ChromaDB知识库检索相关老年医学指南，辅助大模型生成解释。')
add_para('3. 本地大模型：通过Ollama运行qwen2.5:3b模型，生成结构化评估报告，包括Fried 5项逐项核对、阳性计数、结论和个性化临床建议。')
add_para('4. 持续学习机制：收集医生反馈，自动优化模型和Prompt。')

doc.add_page_break()

# 五、评价指标
add_section('五、评价指标')

add_subsection('5.1 主要评价指标')
add_para('人工智能系统与金标准（医生Fried评估）的一致性。')
add_para('评价方法：')
add_para('  - 计算敏感度（Sensitivity）和特异度（Specificity）。')
add_para('  - 计算阳性预测值（PPV）和阴性预测值（NPV）。')
add_para('  - 计算Kappa一致性系数。')
add_para('  - 绘制ROC曲线，计算AUC。')

add_subsection('5.2 次要评价指标')
add_para('1. 人工智能系统与FRAIL量表的一致性（Kappa系数）。')
add_para('2. 人工智能系统的诊断时间（从输入病历到输出结果的时间）。')
add_para('3. 医生诊断时间（从阅读病历到完成Fried评估的时间）。')
add_para('4. 医生使用满意度（问卷调查，Likert 5级量表）。')
add_para('5. 不同衰弱状态（非衰弱、衰弱前期、衰弱）的识别准确率。')

add_subsection('5.3 医生满意度调查')
add_para('采用自编问卷，包含以下维度：')
add_para('  - 系统易用性（5个条目）')
add_para('  - 评估结果准确性（5个条目）')
add_para('  - 临床建议实用性（5个条目）')
add_para('  - 报告可解释性（5个条目）')
add_para('  - 整体满意度（1个条目）')
add_para('每个条目采用Likert 5级评分：1=非常不满意，2=不满意，3=一般，4=满意，5=非常满意。')
add_para('总分21-105分，≥84分为满意。')

doc.add_page_break()

# 六、样本量计算
add_section('六、样本量计算')

add_subsection('6.1 计算依据')
add_para('根据文献报道，老年衰弱患病率约为15%-30%。')
add_para('假设人工智能系统的敏感度为85%，特异度为80%。')
add_para('设定α=0.05（双侧），β=0.20（检验效能80%）。')

add_subsection('6.2 计算方法')
add_para('采用诊断试验样本量计算公式：')
add_para('n = (Zα/2√P(1-P) + Zβ√P1(1-P1) + P0(1-P0))² / (P1-P0)²')
add_para('其中：')
add_para('  - P0 = 预期患病率 = 20%')
add_para('  - P1 = 预期诊断准确率 = 85%')
add_para('  - Zα/2 = 1.96（α=0.05）')
add_para('  - Zβ = 0.84（β=0.20）')

add_subsection('6.3 计算结果')
add_para('经计算，至少需要173例。')
add_para('考虑10%的脱落率，计划纳入200例。')
add_para('其中：')
add_para('  - 非衰弱：约100例（50%）')
add_para('  - 衰弱前期：约60例（30%）')
add_para('  - 衰弱：约40例（20%）')

doc.add_page_break()

# 七、统计分析
add_section('七、统计分析')

add_subsection('7.1 统计软件')
add_para('采用SPSS 26.0或R 4.2进行统计分析。')

add_subsection('7.2 描述性统计')
add_para('计量资料：采用均数±标准差（x̄±s）或中位数（四分位数）描述。')
add_para('计数资料：采用频数（n）和百分比（%）描述。')

add_subsection('7.3 一致性分析')
add_para('1. 人工智能系统 vs 金标准：')
add_para('   - 计算敏感度、特异度、PPV、NPV及其95%置信区间。')
add_para('   - 计算Kappa系数（加权Kappa用于有序分类）。')
add_para('   - Kappa值解释：0.00-0.20=轻微一致，0.21-0.40=一般一致，0.41-0.60=中等一致，0.61-0.80=高度一致，0.81-1.00=几乎完全一致。')
add_para('2. 人工智能系统 vs FRAIL量表：')
add_para('   - 同样计算Kappa系数。')
add_para('3. 2名医生金标准评估的一致性：')
add_para('   - 计算Kappa系数，评估医生间信度。')

add_subsection('7.4 诊断效能比较')
add_para('1. 绘制人工智能系统和FRAIL量表的ROC曲线。')
add_para('2. 计算AUC，比较两者的诊断效能（DeLong检验）。')
add_para('3. 在不同衰弱状态（非衰弱、衰弱前期、衰弱）中分别计算准确率。')

add_subsection('7.5 时间效率比较')
add_para('采用配对t检验或Wilcoxon符号秩检验，比较人工智能系统和医生诊断的时间差异。')

add_subsection('7.6 满意度分析')
add_para('计算各维度得分和总分，采用描述性统计呈现。')
add_para('对满意度与医生年龄、职称、工作年限等进行相关性分析（Pearson或Spearman相关）。')

doc.add_page_break()

# 八、质量控制
add_section('八、质量控制')

add_subsection('8.1 评估者培训')
add_para('1. 对参与金标准评估的医生进行统一培训，确保对Fried标准的理解一致。')
add_para('2. 培训内容包括：Fried 5项指标的定义、测量方法、常见误区。')
add_para('3. 培训后进行预评估，评估10例病历，讨论分歧，统一标准。')

add_subsection('8.2 盲法')
add_para('1. 进行金标准评估的医生不知道人工智能系统的评估结果。')
add_para('2. 人工智能系统的评估自动进行，评估人员不参与。')
add_para('3. 进行FRAIL量表评估的医生不知道金标准结果和人工智能结果。')

add_subsection('8.3 数据管理')
add_para('1. 采用双人双录入，确保数据准确性。')
add_para('2. 建立数据核查机制，定期检查异常值和缺失值。')
add_para('3. 数据存储于加密服务器，仅限研究人员访问。')

doc.add_page_break()

# 九、伦理考虑
add_section('九、伦理考虑')

add_subsection('9.1 伦理审批')
add_para('本研究将向医院伦理委员会提交研究方案，获取伦理审批后方可开展。')

add_subsection('9.2 知情同意')
add_para('1. 向患者或家属详细说明研究目的、方法、风险和获益。')
add_para('2. 签署知情同意书。')
add_para('3. 患者有权随时退出研究，不影响其正常医疗。')

add_subsection('9.3 隐私保护')
add_para('1. 所有病历数据脱敏处理，去除可识别信息。')
add_para('2. 数据仅用于本研究，不用于商业用途。')
add_para('3. 研究结束后，按医院规定销毁原始数据。')
add_para('4. 研究结果发表时，不披露任何可识别患者身份的信息。')

doc.add_page_break()

# 十、论文框架
add_section('十、论文框架（投稿用）')

add_subsection('10.1 题目建议')
add_para('1. 基于人工智能的老年衰弱识别系统构建及临床验证')
add_para('2. 中文电子病历文本的老年衰弱自动识别：一项诊断准确性研究')
add_para('3. BERT-RAG-LLM融合模型在老年衰弱评估中的应用与验证')

add_subsection('10.2 摘要结构（结构化摘要，400-500字）')
add_para('【目的】评估基于人工智能的老年衰弱识别系统对中文电子病历的识别准确性。')
add_para('【方法】收集某医院老年科200例患者的电子病历，由2名高年资医生独立进行Fried衰弱评估作为金标准，人工智能系统和FRAIL量表分别进行评估。计算敏感度、特异度、Kappa系数和AUC，并进行医生满意度调查。')
add_para('【结果】待补充。')
add_para('【结论】待补充。')

add_subsection('10.3 正文结构')
add_para('1. 引言（800-1000字）')
add_para('   - 衰弱定义和流行病学')
add_para('   - 现有评估方法的局限性')
add_para('   - AI在医学文本分析中的应用')
add_para('   - 本研究目的和创新点')
add_para('')
add_para('2. 对象与方法（1500-2000字）')
add_para('   - 研究设计')
add_para('   - 研究对象（纳入/排除标准）')
add_para('   - 金标准评估方法')
add_para('   - 人工智能系统介绍')
add_para('   - FRAIL量表介绍')
add_para('   - 评价指标')
add_para('   - 统计方法')
add_para('')
add_para('3. 结果（1500-2000字）')
add_para('   - 一般资料描述（表1）')
add_para('   - 人工智能系统与金标准的一致性（表2：混淆矩阵）')
add_para('   - 诊断效能指标（表3：敏感度、特异度、Kappa、AUC）')
add_para('   - ROC曲线对比（图1）')
add_para('   - 不同衰弱状态的识别准确率（表4）')
add_para('   - 人工智能系统与FRAIL量表对比（表5）')
add_para('   - 诊断时间对比（表6）')
add_para('   - 医生满意度调查结果（表7）')
add_para('')
add_para('4. 讨论（2000-2500字）')
add_para('   - 主要发现总结')
add_para('   - 与国内外研究对比')
add_para('   - 系统优势分析（可解释性、效率、标准化）')
add_para('   - 局限性（单中心、样本量、模型泛化能力）')
add_para('   - 未来研究方向')
add_para('')
add_para('5. 结论（200-300字）')
add_para('   - 简明总结主要结论')
add_para('   - 临床意义')
add_para('')
add_para('6. 参考文献（20-30篇）')
add_para('   - 近5年文献为主')
add_para('   - 包含本领域经典文献和最新进展')

doc.add_page_break()

# 十一、时间进度
add_section('十一、时间进度（12个月）')

table = doc.add_table(rows=7, cols=3)
table.style = 'Table Grid'
for i, h in enumerate(['阶段', '时间', '内容']):
    cell = table.rows[0].cells[i]
    cell.text = h
    for run in cell.paragraphs[0].runs:
        set_font(run, '宋体', 10, True)
rows = [
    ['第1阶段', '1-3月', '伦理审批、医生培训、病历收集（100例）'],
    ['第2阶段', '4-6月', '病历收集（100例）、金标准评估、系统评估'],
    ['第3阶段', '7-9月', '数据整理、统计分析、FRAIL评估、满意度调查'],
    ['第4阶段', '10-11月', '论文撰写、修改、投稿'],
    ['第5阶段', '12月', '审稿意见回复、修改、录用'],
]
for row_data in rows:
    row = table.add_row()
    for i, text in enumerate(row_data):
        cell = row.cells[i]
        cell.text = text
        for run in cell.paragraphs[0].runs:
            set_font(run, '宋体', 10)

doc.add_page_break()

# 十二、预期成果
add_section('十二、预期成果')

add_subsection('12.1 直接成果')
add_para('1. 论文1篇：投稿《中华老年医学杂志》或《中国老年学杂志》。')
add_para('2. 数据集1个：200例标准化老年衰弱标注数据集。')
add_para('3. 软件著作权1项：老年衰弱智能评估系统。')

add_subsection('12.2 预期结果')
add_para('乐观预期：')
add_para('  - 人工智能系统敏感度≥85%，特异度≥80%')
add_para('  - Kappa系数≥0.75（与金标准高度一致）')
add_para('  - AUC≥0.85')
add_para('  - 诊断时间比医生缩短80%以上')
add_para('  - 医生满意度≥80%')
add_para('')
add_para('保守预期：')
add_para('  - 人工智能系统敏感度≥75%，特异度≥70%')
add_para('  - Kappa系数≥0.60（与金标准中等一致）')
add_para('  - AUC≥0.75')
add_para('  - 诊断时间比医生缩短50%以上')
add_para('  - 医生满意度≥60%')

add_subsection('12.3 后续研究')
add_para('1. 扩大样本量，开展多中心验证。')
add_para('2. 前瞻性研究，验证系统对临床结局（再住院率、死亡率）的预测价值。')
add_para('3. 系统优化后，申请省级/国家级课题。')
add_para('4. 探索系统在其他老年综合征（如肌少症、失能）中的应用。')

# 保存
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '中文核心实验设计方案.docx')
doc.save(output_path)
print(f'已保存: {output_path}')
print(f'文件大小: {os.path.getsize(output_path)} bytes')
