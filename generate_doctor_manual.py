from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

def set_font(run, name='宋体', size=10.5, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color

def add_title(text, level=0):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        set_font(run, '黑体', 22 if level==0 else 16, True, RGBColor(0x2C,0x3E,0x50))
    return p

def add_section(title_text):
    h = doc.add_heading(title_text, level=1)
    for run in h.runs:
        set_font(run, '黑体', 16, True, RGBColor(0x2C,0x3E,0x50))

def add_subsection(title_text):
    h = doc.add_heading(title_text, level=2)
    for run in h.runs:
        set_font(run, '黑体', 13, True, RGBColor(0x2C,0x3E,0x50))

def add_para(text, bold=False, color=None):
    p = doc.add_paragraph(text)
    for run in p.runs:
        set_font(run, '宋体', 11, bold)
        if color:
            run.font.color.rgb = color
    return p

def add_tip(text):
    p = doc.add_paragraph()
    set_font(p.add_run('💡 提示：'), '宋体', 11, True, RGBColor(0x27,0xAE,0x60))
    set_font(p.add_run(text), '宋体', 11)

def add_warning(text):
    p = doc.add_paragraph()
    set_font(p.add_run('⚠️ 注意：'), '宋体', 11, True, RGBColor(0xE7,0x4C,0x3C))
    set_font(p.add_run(text), '宋体', 11)

# 封面
add_title('老年衰弱智能评估系统')
add_title('医生操作手册', level=1)
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run('—— 非技术人员专用，无需编程，全程鼠标操作 ——'), '楷体', 12, False, RGBColor(0x7F,0x8C,0x8D))
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run('版本：v2.0  |  更新日期：2026-07-20'), '宋体', 10, False, RGBColor(0x95,0xA5,0xA6))

doc.add_page_break()

# 目录
add_section('目录')
items = [
    '一、开始之前：你需要准备什么',
    '二、每天第一步：启动系统',
    '三、给单个病人做评估',
    '四、批量导入病历（一次处理很多病人）',
    '五、批量标注与训练模型（进阶功能）',
    '六、查看和导出数据',
    '七、常见问题：遇到问题怎么办',
    '八、重要提醒',
]
for item in items:
    p = doc.add_paragraph(item, style='List Number')
    for run in p.runs:
        set_font(run, '宋体', 11)

doc.add_page_break()

# 一、开始之前
add_section('一、开始之前：你需要准备什么')

add_para('这份手册是写给医生的，假设你：')
add_para('• 会用电脑打开文件、复制粘贴文字')
add_para('• 不会写代码，也不想学代码')
add_para('• 想要用这套系统来评估老年病人是否衰弱')
add_para('• 想要积累数据，训练出更适合你们医院的模型')
add_para('')
add_para('在开始使用之前，请确认以下准备工作已经完成：')

add_subsection('1.1 电脑要求')
add_para('• 操作系统：Windows 10 或 Windows 11')
add_para('• 内存：至少 8GB（建议 16GB）')
add_para('• 硬盘：至少 50GB 剩余空间')
add_para('• 显卡：有独立显卡最好，没有也能运行（只是慢一些）')

add_subsection('1.2 软件安装')
add_para('以下软件由开发团队帮你安装好，你不需要自己操作：')
add_para('• Python 程序环境')
add_para('• Ollama（运行大模型的软件）')
add_para('• BERT 模型文件（约 400MB）')
add_para('• BGE 向量模型（约 1.3GB）')
add_para('')
add_warning('这些安装工作由开发团队完成，你不需要自己动手。如果电脑重装系统或换电脑，请联系开发团队重新安装。')

add_subsection('1.3 病历数据格式')
add_para('如果你要批量导入病历，请让医院信息科导出 CSV 或 Excel 文件，要求：')
add_para('• 文件格式：.csv 或 .xlsx')
add_para('• 必须包含一列"病历文本"（列名可以是：病历文本、text、入院记录、病情描述）')
add_para('• 建议包含：患者ID、姓名、年龄、性别')
add_para('• 病历内容需要脱敏：删除真实姓名、身份证号、手机号等敏感信息，用"患者A"、"张三"等代替')
add_para('')
add_tip('如果不确定信息科导出的格式对不对，可以先导出一小份（比如10条），试导入看看能否成功。')

doc.add_page_break()

# 二、每天第一步：启动系统
add_section('二、每天第一步：启动系统')

add_subsection('2.1 启动 Ollama（大模型服务）')
add_para('Ollama 是运行人工智能大模型的后台程序，必须先启动它，系统才能正常工作。')
add_para('')
add_para('操作步骤：')
add_para('1. 在电脑桌面左下角点击"开始"按钮')
add_para('2. 搜索"cmd"或"命令提示符"，点击打开黑色窗口')
add_para('3. 在黑色窗口中输入以下命令，然后按回车键：')
add_para('')
p = doc.add_paragraph()
p.paragraph_format.left_indent = Inches(0.5)
set_font(p.add_run('ollama serve'), 'Consolas', 11, True, RGBColor(0x2C,0x3E,0x50))
add_para('')
add_para('4. 看到显示"Listening on 127.0.0.1:11434"就表示启动成功了')
add_para('5. 这个黑色窗口不要关闭，最小化到任务栏即可')
add_para('')
add_tip('可以把这个黑色窗口最小化，但不要关闭它。关闭后系统就无法调用大模型了。')
add_warning('如果提示"ollama 不是内部或外部命令"，说明 Ollama 没有安装好，请联系开发团队。')

add_subsection('2.2 启动主程序')
add_para('操作步骤：')
add_para('1. 打开"我的电脑"或"文件资源管理器"')
add_para('2. 找到系统所在的文件夹（开发团队会告诉你具体路径，通常是 D 盘某个文件夹）')
add_para('3. 依次打开：pythonProject → source')
add_para('4. 找到名为 frailty_assessment_app.py 的文件')
add_para('5. 在这个文件上点击鼠标右键 → 选择"打开方式" → 选择"Python"（或 python3）')
add_para('6. 等待几秒钟，会出现一个窗口，显示"系统就绪"，表示启动成功')
add_para('')
add_tip('第一次启动会比较慢（约1-2分钟），因为需要加载模型文件。以后启动会快很多。')
add_para('')
add_warning('如果启动时弹出错误窗口，请截图保存，联系开发团队解决。常见原因是 Ollama 没有启动。')

doc.add_page_break()

# 三、给单个病人做评估
add_section('三、给单个病人做评估')

add_para('这是系统最常用的功能：输入一个病人的病历，系统自动判断他是不是衰弱。')

add_subsection('3.1 输入病历')
add_para('1. 在系统左侧的大文本框中，粘贴或输入病人的病历文本')
add_para('2. 病历内容应该包括：主诉、现病史、既往史、查体结果（体重、握力、步速、活动量等）')
add_para('3. 如果不确定写什么，可以点击左下角的【示例】按钮，看看示例病历的格式')
add_para('')
add_tip('病历写得越详细，系统判断越准确。特别是体重变化、握力、步速、活动量这些数据很重要。')

add_subsection('3.2 快速预测（BERT 自动分类）')
add_para('1. 点击上方工具栏的【快速预测】按钮（或按菜单【工具】→【快速预测】）')
add_para('2. 等待约 1-3 秒')
add_para('3. 在右上方的"快速预测"标签页中，看到结果：')
add_para('   • 预测结论：衰弱 / 非衰弱')
add_para('   • 置信度：0-1 之间的数字，越接近 1 表示系统越确定')
add_para('')
add_tip('置信度低于 0.7 时，系统不太确定，建议再仔细看一遍病历，或者做详细解释。')

add_subsection('3.3 详细解释（AI 生成完整报告）')
add_para('1. 点击上方工具栏的【详细解释】按钮')
add_para('2. 等待约 10-30 秒（系统正在调用大模型分析）')
add_para('3. 在右下方的"详细解释"标签页中，看到完整的结构化报告：')
add_para('   • 第1部分：Fried 衰弱表型逐项核对（5项指标，每项显示阳性/阴性）')
add_para('   • 第2部分：阳性指标总数（0-5 项）')
add_para('   • 第3部分：最终结论（衰弱 / 衰弱前期 / 非衰弱）')
add_para('   • 第4部分：临床建议')
add_para('')
add_para('报告中的每一项都有"证据"引用，告诉你系统是根据病历中哪句话判断的。')
add_para('')
add_tip('如果报告中说"未提及"，说明病历里缺少这项指标的信息。你可以补充后再重新评估。')

add_subsection('3.4 如果结果不对，提交反馈')
add_para('如果你发现系统的判断和临床实际不一致：')
add_para('1. 做完【详细解释】后，点击右下角的【💡 提交反馈】按钮')
add_para('2. 在弹出的窗口中：')
add_para('   • 选择"你认为正确的结论"（衰弱 / 衰弱前期 / 非衰弱）')
add_para('   • 填写"校正原因"（比如："患者是骨折后卧床，不是衰弱"）')
add_para('   • 填写你的姓名（方便后续统计）')
add_para('3. 点击【确定】')
add_para('')
add_para('你的反馈会自动保存，系统会用它来改进模型。')
add_para('')
add_tip('反馈越多，系统越聪明。即使你觉得系统判断对了，也可以不反馈。只有判断错了才需要反馈。')
add_warning('反馈数据保存在本地电脑，不会上传到互联网，不用担心隐私泄露。')

doc.add_page_break()

# 四、批量导入病历
add_section('四、批量导入病历（一次处理很多病人）')

add_para('如果你有一批病人的病历（比如从信息科导出的表格），可以一次性导入，系统自动评估所有病人。')

add_subsection('4.1 准备文件')
add_para('请医院信息科导出 CSV 或 Excel 文件，要求：')
add_para('• 每行一个病人')
add_para('• 必须有一列是病历文本（列名可以是：病历文本、text、入院记录、病情描述）')
add_para('• 病历内容已脱敏（去除真实姓名、身份证号等）')
add_para('')
add_tip('如果信息科导出的格式不对，可以把文件发给我（开发团队），我帮你转换。')

add_subsection('4.2 导入操作')
add_para('1. 点击菜单【文件】→【批量导入...】')
add_para('2. 在弹出的文件选择窗口中，找到你的 CSV/Excel 文件')
add_para('3. 可以一次选择多个文件')
add_para('4. 点击【打开】')
add_para('5. 系统开始自动处理，会显示进度条')
add_para('6. 处理完成后，中间的患者列表表格会显示所有病人的评估结果')
add_para('')
add_para('表格中：')
add_para('• 红色姓名 = 衰弱')
add_para('• 橙色姓名 = 衰弱前期')
add_para('• 绿色姓名 = 非衰弱')
add_para('• 置信度 = 系统判断的确定程度')

add_subsection('4.3 查看单个病人详情')
add_para('在表格中双击任意一行，系统会把该病人的病历加载到左侧输入区，你可以：')
add_para('• 查看完整病历文本')
add_para('• 重新做详细解释')
add_para('• 提交反馈')

add_subsection('4.4 导出结果')
add_para('1. 点击中间表格上方的【导出 Excel】按钮')
add_para('2. 选择保存位置')
add_para('3. 系统会生成一个 Excel 文件，包含所有病人的评估结果')
add_para('')
add_tip('导出的 Excel 可以用于科室汇报、统计分析、或者作为科研数据。')

doc.add_page_break()

# 五、批量标注与训练模型
add_section('五、批量标注与训练模型（进阶功能）')

add_para('这个功能用于：')
add_para('• 你有大量历史病历，想要系统帮你预标注，然后你审核修正')
add_para('• 积累了一定数量的标注数据后，训练一个更适合你们医院的模型')
add_para('')
add_warning('这个功能需要你对"衰弱"的判断有比较准确的把握。如果不确定，建议先多使用"单个评估"功能，熟悉系统后再操作。')

add_subsection('5.1 打开批量标注窗口')
add_para('1. 点击菜单【工具】→【批量标注与模型训练...】')
add_para('2. 会弹出一个新窗口，标题是"批量标注与模型训练"')
add_para('3. 这个窗口独立于主窗口，可以同时开着')

add_subsection('5.2 导入脱密病历')
add_para('1. 在新窗口中，找到"步骤1：导入脱密病历"')
add_para('2. 点击【选择 CSV/Excel 文件】按钮')
add_para('3. 选择你的病历文件')
add_para('4. 系统会自动做 BERT 预标注，置信度低的样本会标记为"不确定"')
add_para('5. 等待导入完成，状态栏会显示"导入完成：新增 XX 条记录"')

add_subsection('5.3 LLM 预标注（可选）')
add_para('1. 点击【运行 LLM 预标注】按钮')
add_para('2. 系统会对未标注的记录调用大模型，生成详细评估')
add_para('3. 这个过程比较慢，可能需要几分钟到几十分钟，取决于数据量')
add_para('4. 你可以最小化窗口，去做其他事情')
add_para('')
add_tip('LLM 预标注不是必须的。如果数据量很大，可以跳过这一步，直接做人工审核。')

add_subsection('5.4 人工审核（核心步骤）')
add_para('1. 点击【加载待审核记录】按钮')
add_para('2. 左侧表格会显示待审核的记录')
add_para('3. 黄色高亮的行 = 系统不确定的样本（优先审核）')
add_para('4. 点击表格中的任意一行：')
add_para('   • 右侧会显示该病人的完整病历原文')
add_para('   • 你可以仔细阅读病历，判断衰弱状态')
add_para('5. 在"审核人"输入框中填写你的姓名')
add_para('6. 点击三个按钮之一：')
add_para('   • 【衰弱】- 红色按钮')
add_para('   • 【衰弱前期】- 橙色按钮')
add_para('   • 【非衰弱】- 绿色按钮')
add_para('7. 已标注的记录会从表格中消失，继续审核下一条')
add_para('')
add_tip('建议优先审核黄色高亮的不确定样本，这些样本对模型训练最有价值。')
add_warning('审核时请认真判断，你的标注会作为"正确答案"用来训练模型。如果标注错误，模型会学到错误知识。')

add_subsection('5.5 查看标注质量')
add_para('1. 点击【查看标注质量分析】按钮')
add_para('2. 系统会显示统计报告：')
add_para('   • 总记录数')
add_para('   • 已人工标注的比例')
add_para('   • BERT 不确定的样本数')
add_para('   • BERT 和 LLM 结论一致的比例')
add_para('')
add_tip('BERT-LLM 一致性低的样本，说明系统内部有分歧，建议重点审核这些样本。')

add_subsection('5.6 训练模型')
add_para('当你积累了足够多的标注数据（建议至少 100 条），可以训练新模型：')
add_para('1. 在"步骤6：模型重训练"区域：')
add_para('   • "最少样本"默认是 100，可以改')
add_para('   • "Epoch"默认是 5，可以改（数字越大训练时间越长，一般 5-10 即可）')
add_para('2. 点击【全量重训练 BERT】按钮')
add_para('3. 确认提示后，系统开始训练')
add_para('4. 训练时间取决于数据量和电脑性能，可能需要 10 分钟到 1 小时')
add_para('5. 训练完成后，系统会提示"是否切换到新模型"')
add_para('6. 点击【是】，系统立即使用新模型（不需要重启程序）')
add_para('')
add_tip('第一次训练建议用默认参数。有经验后可以调整参数尝试更好的效果。')
add_warning('训练过程中电脑会比较卡，建议不要同时做其他事情。训练时可以去做别的事，让电脑自己运行。')

add_subsection('5.7 导出标注数据')
add_para('1. 点击【导出标注数据（CSV）】按钮')
add_para('2. 选择保存位置')
add_para('3. 系统会导出所有已人工审核的标注记录')
add_para('')
add_tip('导出的 CSV 可以用于：统计分析、论文写作、或者发给开发团队做进一步分析。')

doc.add_page_break()

# 六、查看和导出数据
add_section('六、查看和导出数据')

add_subsection('6.1 反馈数据在哪里')
add_para('你提交的反馈保存在以下位置（开发团队会告诉你具体路径）：')
add_para('• feedback.db - 数据库文件')
add_para('• feedback_log.jsonl - 日志文件（双保险）')
add_para('')
add_para('这些文件可以用 Excel 打开查看，或者发给开发团队分析。')

add_subsection('6.2 批量标注数据在哪里')
add_para('• batch_annotation.db - 批量标注数据库')
add_para('• 可以通过"导出标注数据"按钮导出为 CSV')

add_subsection('6.3 训练好的模型在哪里')
add_para('• frailty_bert_retrained_v2/ - 第2版模型')
add_para('• frailty_bert_retrained_v3/ - 第3版模型')
add_para('• 以此类推')
add_para('')
add_para('系统会自动使用最新版本的模型。')

doc.add_page_break()

# 七、常见问题
add_section('七、常见问题：遇到问题怎么办')

add_subsection('Q1：启动时黑窗口闪一下就没了')
add_para('可能原因：Ollama 没有启动，或者模型文件损坏。')
add_para('解决方法：')
add_para('1. 先确认 Ollama 已经启动（见第二章 2.1）')
add_para('2. 如果还是不行，截图错误信息，发给开发团队')

add_subsection('Q2：详细解释一直显示"等待大模型分析中"')
add_para('可能原因：Ollama 服务卡住了，或者电脑性能不足。')
add_para('解决方法：')
add_para('1. 等待 1-2 分钟，看是否会自动完成')
add_para('2. 如果一直卡住，关闭系统，重新启动 Ollama，再重新打开系统')
add_para('3. 如果还是不行，联系开发团队')

add_subsection('Q3：弹出"病历数据缺失警告"')
add_para('这不是错误，是系统提醒你：病历里缺少一些关键指标。')
add_para('你可以：')
add_para('• 点击【否】- 去补充病历信息后再评估')
add_para('• 点击【是】- 继续评估（结果可能不太准）')
add_para('')
add_tip('如果你确定病历已经写得很全了，但系统还是提示缺失，可能是病历里用了不同的词描述。可以忽略警告，继续评估。')

add_subsection('Q4：批量导入提示"未找到病历文本列"')
add_para('可能原因：CSV/Excel 文件的列名不对。')
add_para('解决方法：')
add_para('1. 用 Excel 打开文件，看看列名是什么')
add_para('2. 把病历文本那一列的列名改成"病历文本"或"text"')
add_para('3. 保存后重新导入')
add_para('4. 如果还是不行，把文件发给开发团队帮忙处理')

add_subsection('Q5：训练模型提示"样本不足"')
add_para('原因：你标注的数据还不够 100 条（默认阈值）。')
add_para('解决方法：')
add_para('1. 继续人工审核，积累更多标注数据')
add_para('2. 或者把"最少样本"改成更小的数字（比如 50），但不建议太小，否则模型效果不好')

add_subsection('Q6：系统突然变得很卡')
add_para('可能原因：')
add_para('• 正在训练模型（CPU/GPU 占用高）')
add_para('• 同时打开了太多程序')
add_para('• 电脑内存不足')
add_para('解决方法：')
add_para('1. 如果正在训练，耐心等待')
add_para('2. 关闭其他不用的程序')
add_para('3. 重启电脑')

add_subsection('Q7：我想清理测试数据，重新开始')
add_para('联系开发团队，或者按以下步骤操作：')
add_para('1. 关闭系统')
add_para('2. 找到系统文件夹，删除以下文件：')
add_para('   • batch_annotation.db')
add_para('   • feedback.db')
add_para('   • feedback_log.jsonl')
add_para('   • frailty_bert_retrained_v*（所有以这个开头的文件夹）')
add_para('   • frailty_bert_retrain（文件夹）')
add_para('3. 重新启动系统，数据就清空了')
add_para('')
add_warning('删除前请确认这些数据确实不需要了！删除后无法恢复。')

doc.add_page_break()

# 八、重要提醒
add_section('八、重要提醒')

add_subsection('8.1 数据安全')
add_para('• 所有数据保存在你的本地电脑，不会自动上传到互联网')
add_para('• 导入的病历必须已经脱敏（去除真实姓名、身份证号、手机号等）')
add_para('• 定期备份 feedback.db 和 batch_annotation.db 文件')
add_para('• 不要把系统文件夹复制到公共电脑或U盘上')

add_subsection('8.2 系统局限性')
add_para('• 系统目前使用的是 3B 小模型，能力有限，复杂病例可能判断错误')
add_para('• BERT 模型只在演示数据上训练过，在你们医院的真实数据上可能准确率不够高')
add_para('• 系统的输出仅供参考，不能替代医生的临床判断')
add_para('• 最终诊断必须由医生确认')

add_subsection('8.3 使用建议')
add_para('• 初期以"单个评估"为主，熟悉系统后再用批量功能')
add_para('• 多提交反馈，帮助系统改进')
add_para('• 积累了足够数据后，再训练模型')
add_para('• 训练模型时先用默认参数，有经验后再调整')
add_para('• 定期导出数据，做备份')

add_subsection('8.4 联系方式')
add_para('遇到任何问题，请：')
add_para('1. 先查看本手册的"常见问题"章节')
add_para('2. 如果手册解决不了，截图错误信息')
add_para('3. 联系开发团队（微信/电话/邮件）')
add_para('')
add_para('截图方法：')
add_para('• 按键盘上的 PrtScn 键，然后粘贴到微信聊天窗口')
add_para('• 或者按 Win + Shift + S，选择要截图的区域')

add_subsection('8.5 最后的话')
add_para('这套系统的目标是成为你的辅助工具，帮你更快、更标准地识别老年衰弱。')
add_para('但它不是完美的，需要你的反馈来不断改进。')
add_para('每一次你提交的反馈，都是在教系统变得更聪明。')
add_para('感谢你的耐心使用和宝贵意见！')

# 保存
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), '医生操作手册_老年衰弱智能评估系统.docx')
doc.save(output_path)
print(f'已保存: {output_path}')
print(f'文件大小: {os.path.getsize(output_path)} bytes')
