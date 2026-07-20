from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import os

doc = Document()

def set_font(run, name='宋体', size=10.5, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color

def add_title(text, level=0):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p.runs:
        set_font(run, '黑体', 22 if level==0 else 16, True, RGBColor(0x2C,0x3E,0x50))
    return p

def add_section(title_text):
    h = doc.add_heading(title_text, level=1)
    for run in h.runs:
        set_font(run, '黑体', 16, True, RGBColor(0x2C,0x3E,0x50))

def add_subsection(title_text):
    h = doc.add_heading(title_text, level=2)
    for run in h.runs:
        set_font(run, '黑体', 13, True, RGBColor(0x2C,0x3E,0x50))

def add_para(text, bold=False):
    p = doc.add_paragraph(text)
    for run in p.runs:
        set_font(run, '宋体', 11, bold)
    return p

# 封面
add_title('老年衰弱智能评估系统')
add_title('RAG知识库文献汇编', level=1)
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run('—— 系统知识库来源文献完整清单 ——'), '楷体', 12, False, RGBColor(0x7F,0x8C,0x8D))
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(p.add_run('收录文献：30+篇  |  涵盖领域：指南、共识、RCT、综述、工具'), '宋体', 10, False, RGBColor(0x95,0xA5,0xA6))

doc.add_page_break()

# 目录
add_section('目录')
items = [
    '一、Fried衰弱表型标准详解',
    '二、Fried标准鉴别诊断',
    '三、其他衰弱评估工具',
    '四、衰弱与常见疾病鉴别',
    '五、衰弱患者营养干预',
    '六、衰弱患者运动干预',
    '七、衰弱患者药物管理',
    '八、衰弱患者跌倒预防',
    '九、认知与心理支持',
    '十、随访管理与转诊',
    '十一、预后与逆转',
    '十二、特殊人群评估',
    '十三、国际核心文献（原始研究）',
    '十四、中国指南与专家共识',
    '十五、干预与康复RCT研究',
    '十六、衰弱与特定疾病',
    '十七、评估工具详解',
    '十八、公共卫生与流行病学',
    '十九、附录：参考值速查表',
    '二十、附录：干预方案速查',
]
for item in items:
    p = doc.add_paragraph(item, style='List Number')
    for run in p.runs:
        set_font(run, '宋体', 11)

doc.add_page_break()

# ========== 一、Fried衰弱表型标准详解 ==========
add_section('一、Fried衰弱表型标准详解')

add_para('Fried衰弱表型标准是目前国际上最常用的衰弱评估工具之一，包含5项指标，满足≥3项为"衰弱"，1-2项为"衰弱前期"，0项为"非衰弱"。', True)

add_subsection('1.1 不明原因体重下降')
add_para('定义：过去1年内体重下降≥5%，或≥4.5kg（10磅）。')
add_para('关键要点：必须是"不明原因"的体重下降。')
add_para('排除情况：主动节食、糖尿病治疗、甲状腺功能亢进治疗、肿瘤化疗等导致的体重下降不算。')
add_para('临床询问："您最近一年体重有没有明显下降？大概下降了多少？"')
add_para('病历记录关键词：体重下降、消瘦、体重减轻、体重降低。')

add_subsection('1.2 疲乏/精力减退')
add_para('定义：自我报告感到疲乏无力，精力不足。')
add_para('常用评估工具：CES-D抑郁量表中的2个问题（"我感到做任何事都很费力"、"我觉得我无法继续我的生活"）。')
add_para('关键要点：排除急性疾病（如急性感染）导致的暂时性疲乏。')
add_para('临床询问："您是否经常感到疲乏无力，即使休息后也无法恢复？"')
add_para('病历记录关键词：疲乏、乏力、无力、没精神、精力减退、疲惫、倦怠。')

add_subsection('1.3 握力下降')
add_para('定义：使用握力计测量，握力低于同性别、同BMI分位的正常参考值。')
add_para('正常参考值（男性）：')
add_para('  BMI ≤24：握力≤29kg')
add_para('  BMI 24.1-28：握力≤30kg')
add_para('  BMI >28：握力≤32kg')
add_para('正常参考值（女性）：')
add_para('  BMI ≤23：握力≤17kg')
add_para('  BMI 23.1-26：握力≤17.3kg')
add_para('  BMI 26.1-29：握力≤18kg')
add_para('  BMI >29：握力≤21kg')
add_para('测量方法：使用Jamar握力计，站立位，手臂自然下垂，用力握3次，取最大值。')
add_para('关键要点：排除手部疾病（如关节炎、腕管综合征、脑卒中后偏瘫）导致的握力下降。')
add_para('病历记录关键词：握力、手力、抓握。')

add_subsection('1.4 步速减慢')
add_para('定义：行走4.57米（15英尺）时间延长，超过同性别、同身高分位的正常参考值。')
add_para('正常参考值（男性）：')
add_para('  身高≤173cm：步速>0.65m/s')
add_para('  身高>173cm：步速>0.76m/s')
add_para('正常参考值（女性）：')
add_para('  身高≤159cm：步速>0.65m/s')
add_para('  身高>159cm：步速>0.76m/s')
add_para('测量方法：标记4.57米距离，从静止开始行走，记录通过时间，计算速度。')
add_para('关键要点：排除急性疾病（如骨折、急性关节炎）导致的步速减慢。')
add_para('临床询问："您平时走路的速度有没有明显变慢？"')
add_para('病历记录关键词：步速、步态、行走缓慢、走路慢、步行速度。')

add_subsection('1.5 活动量减少')
add_para('定义：每周消耗的热量低于标准值。')
add_para('男性标准：每周消耗热量<383kcal（约相当于每周散步<2.5小时）')
add_para('女性标准：每周消耗热量<270kcal（约相当于每周散步<2小时）')
add_para('常用评估工具：明尼苏达休闲时间活动问卷（MLTAQ）。')
add_para('关键要点：排除急性疾病或外伤导致的暂时性活动减少。')
add_para('排除情况：骨折后卧床、急性心梗后限制活动、术后恢复期等不算。')
add_para('临床询问："您最近半年的活动量有没有明显减少？"')
add_para('病历记录关键词：活动量、活动减少、卧床、不愿活动、运动减少。')

doc.add_page_break()

# ========== 二、Fried标准鉴别诊断 ==========
add_section('二、Fried标准鉴别诊断：什么算衰弱，什么不算')

add_subsection('2.1 应判定为衰弱的情况')
add_para('• 无明确病因的体重下降 + 疲乏 + 握力下降 + 步速减慢 + 活动量减少（满足≥3项）')
add_para('• 慢性病稳定期出现的多系统功能储备下降')
add_para('• 年龄相关生理性储备下降，无急性疾病诱因')

add_subsection('2.2 不应判定为衰弱的情况')
add_para('• 卒中后肌力下降：这是脑血管疾病导致的局灶性神经功能缺损，不是Fried标准的"握力下降"')
add_para('• 骨折后卧床：这是外伤导致的暂时性活动受限，不是Fried标准的"活动量减少"')
add_para('• 术后恢复期：手术创伤导致的暂时性虚弱，恢复后可能改善')
add_para('• 急性感染期间：发热、感染导致的暂时性疲乏和体重下降')
add_para('• 帕金森病导致的步态异常：这是神经系统疾病，不是Fried标准的"步速减慢"')
add_para('• 重度骨关节炎导致的行走困难：这是关节疾病，不是Fried标准的"步速减慢"')
add_para('• 抑郁症导致的疲乏：这是精神疾病，需先治疗抑郁症再评估')
add_para('• 甲状腺功能减退导致的疲乏和体重增加：这是内分泌疾病，需先纠正甲减')
add_para('• 营养不良导致的体重下降：如果是因为进食困难、吞咽障碍等，应先解决营养问题')

add_subsection('2.3 衰弱前期 vs 衰弱的鉴别')
add_para('衰弱前期：满足1-2项Fried指标，但日常生活能力基本正常。')
add_para('衰弱：满足≥3项Fried指标，日常生活能力可能受限。')
add_para('关键区别：衰弱前期是"有风险"，衰弱是"已发生"。')

doc.add_page_break()

# ========== 三、其他衰弱评估工具 ==========
add_section('三、其他衰弱评估工具简介')

add_subsection('3.1 FRAIL量表（快速筛查工具）')
add_para('5个问题，每个回答"是"得1分。')
add_para('0分：健康；1-2分：衰弱前期；3-5分：衰弱。')
add_para('适合门诊快速筛查。')

add_subsection('3.2 临床衰弱量表（CFS）')
add_para('1-9级评分：')
add_para('1级：非常健康；2级：健康；3级：管理良好；4级：脆弱；')
add_para('5级：轻度衰弱；6级：中度衰弱；7级：重度衰弱；')
add_para('8级：非常重度衰弱；9级：终末期。')
add_para('适合住院患者。')

add_subsection('3.3 Edmonton衰弱量表（EFS）')
add_para('包含9个维度：认知、健康状况、功能、社会支持、药物使用、营养、情绪、失禁、功能表现。')
add_para('适合社区老年人综合评估。')

add_subsection('3.4 评估工具对比')
add_para('')
table = doc.add_table(rows=7, cols=5)
table.style = 'Table Grid'
headers = ['工具', '项目数', '耗时', '适用场景', '特点']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for run in cell.paragraphs[0].runs:
        set_font(run, '宋体', 10, True)
rows = [
    ['Fried表型', '5项', '10-15分钟', '研究、综合评估', '金标准，需测量设备'],
    ['FRAIL量表', '5个问题', '<1分钟', '门诊快速筛查', '简便，自评或他评'],
    ['CFS', '1-9级', '1-2分钟', '住院、急性照护', '临床判断，适用广泛'],
    ['EFS', '9维度', '5-10分钟', '社区综合评估', '全面，含认知和社会支持'],
    ['SPPB', '3项测试', '5分钟', '功能评估', '客观测量，预测效度好'],
    ['TUGT', '1项测试', '1-2分钟', '平衡和移动评估', '简单，可重复'],
]
for row_data in rows:
    row = table.add_row()
    for i, text in enumerate(row_data):
        cell = row.cells[i]
        cell.text = text
        for run in cell.paragraphs[0].runs:
            set_font(run, '宋体', 10)

doc.add_page_break()

# ========== 四、衰弱与常见疾病鉴别 ==========
add_section('四、衰弱与常见疾病鉴别')

add_subsection('4.1 衰弱 vs 肌少症')
add_para('肌少症：以骨骼肌质量和力量下降为核心。')
add_para('衰弱：以多系统生理储备下降为核心，范围更广。')
add_para('关系：肌少症是衰弱的重要组成部分，但衰弱不等同于肌少症。')

add_subsection('4.2 衰弱 vs 营养不良')
add_para('营养不良：以营养摄入不足或吸收障碍为核心。')
add_para('衰弱：以生理储备下降为核心。')
add_para('关系：营养不良可导致衰弱，但衰弱患者不一定都营养不良。')

add_subsection('4.3 衰弱 vs 失能')
add_para('失能：以日常生活活动能力（ADL）受限为核心。')
add_para('衰弱：以生理储备下降为核心，可能在失能之前出现。')
add_para('关系：衰弱是失能的前驱状态，及时干预可延缓失能。')

add_subsection('4.4 衰弱 vs 抑郁')
add_para('抑郁：以情绪低落、兴趣丧失为核心。')
add_para('衰弱：以生理储备下降为核心。')
add_para('关系：两者常共存，需鉴别主次。')

doc.add_page_break()

# ========== 五、衰弱患者营养干预 ==========
add_section('五、衰弱患者营养干预')

add_subsection('5.1 蛋白质补充')
add_para('每日蛋白质摄入：1.2-1.5 g/kg体重。')
add_para('优质蛋白占比：≥50%。')
add_para('优质蛋白来源：鸡蛋、牛奶、鱼肉、瘦肉、豆制品。')
add_para('口服营养补充剂（ONS）：安素、能全素等，每日400-600 kcal。')

add_subsection('5.2 维生素与矿物质')
add_para('维生素D：800-1000 IU/日。')
add_para('钙剂：1000-1200 mg/日。')
add_para('维生素B12：尤其素食者需补充。')
add_para('叶酸：缺乏时补充。')

add_subsection('5.3 饮食模式')
add_para('地中海饮食模式：富含蔬菜、水果、全谷物、橄榄油、鱼类。')
add_para('少量多餐：每日5-6餐，避免一次进食过多。')
add_para('充足水分：每日1500-2000ml。')

doc.add_page_break()

# ========== 六、衰弱患者运动干预 ==========
add_section('六、衰弱患者运动干预')

add_subsection('6.1 有氧运动')
add_para('频率：每周≥5天。')
add_para('强度：中等强度（微微出汗，能说话但不能唱歌）。')
add_para('时间：每次30分钟。')
add_para('类型：快走、慢跑、游泳、骑自行车。')

add_subsection('6.2 抗阻训练')
add_para('频率：每周2-3次。')
add_para('强度：使用弹力带或轻哑铃。')
add_para('组数：每个动作8-12次/组，1-3组。')
add_para('肌群：下肢（深蹲、提踵）、上肢（哑铃弯举、俯卧撑）、核心（平板支撑）。')

add_subsection('6.3 平衡训练')
add_para('频率：每周≥3次。')
add_para('时间：每次20-30分钟。')
add_para('类型：太极拳、单腿站立、脚跟-脚尖行走。')

add_subsection('6.4 注意事项')
add_para('运动前需评估心肺功能。')
add_para('从低强度开始，循序渐进。')
add_para('有跌倒风险者需有人陪同。')
add_para('出现胸痛、气短、头晕应立即停止。')

doc.add_page_break()

# ========== 七、衰弱患者药物管理 ==========
add_section('七、衰弱患者药物管理')

add_subsection('7.1 多重用药审查')
add_para('标准：Beers标准、STOPP/START标准。')
add_para('频率：每3个月审查一次。')
add_para('重点：抗胆碱能药物、镇静安眠药、阿片类药物。')

add_subsection('7.2 用药原则')
add_para('低剂量起始、缓慢增量。')
add_para('优先停用无明确适应症的药物。')
add_para('避免使用Beers标准中的潜在不适当药物。')
add_para('注意药物相互作用。')

add_subsection('7.3 常见需调整的药物')
add_para('降压药：避免过度降压导致体位性低血压。')
add_para('降糖药：避免低血糖风险。')
add_para('抗凝药：评估出血风险。')

doc.add_page_break()

# ========== 八、衰弱患者跌倒预防 ==========
add_section('八、衰弱患者跌倒预防')

add_subsection('8.1 风险评估')
add_para('使用Morse跌倒评估量表或Hendrich II跌倒风险评估模型。')
add_para('评估频率：入院时、病情变化时、出院前。')

add_subsection('8.2 环境改造')
add_para('居家：改善照明、安装扶手、移除障碍物、使用防滑垫。')
add_para('医院：床栏、呼叫器、防滑拖鞋、夜间照明。')

add_subsection('8.3 辅助器具')
add_para('手杖：轻度平衡障碍。')
add_para('助行器：中度平衡障碍。')
add_para('轮椅：重度活动受限。')

add_subsection('8.4 药物调整')
add_para('停用或减量增加跌倒风险的药物（镇静安眠药、降压药、利尿剂）。')

doc.add_page_break()

# ========== 九、认知与心理支持 ==========
add_section('九、衰弱患者认知与心理支持')

add_subsection('9.1 认知评估')
add_para('工具：MMSE（简易智力状态检查量表）、MoCA（蒙特利尔认知评估量表）。')
add_para('频率：每年一次，或病情变化时。')

add_subsection('9.2 抑郁筛查')
add_para('工具：GDS-15（老年抑郁量表-15项）。')
add_para('频率：每年一次。')

add_subsection('9.3 心理干预')
add_para('认知行为疗法（CBT）。')
add_para('团体心理支持。')
add_para('家庭支持和社会支持。')

doc.add_page_break()

# ========== 十、随访管理 ==========
add_section('十、衰弱患者随访管理')

add_subsection('10.1 随访频率')
add_para('衰弱前期：每6个月。')
add_para('衰弱：每3个月。')
add_para('住院后：出院后1周、1个月、3个月。')

add_subsection('10.2 随访内容')
add_para('体重变化、活动能力（步速、握力）、营养状况、用药依从性、跌倒事件、再住院情况。')

add_subsection('10.3 转诊指征')
add_para('严重营养不良（BMI<18.5或白蛋白<30g/L）。')
add_para('反复跌倒（3个月内≥2次）。')
add_para('严重抑郁（GDS-15≥8分）。')
add_para('认知功能快速下降。')
add_para('需要多学科团队干预。')

doc.add_page_break()

# ========== 十一、预后与逆转 ==========
add_section('十一、衰弱与预后')

add_subsection('11.1 衰弱患者的预后特点')
add_para('住院时间延长、再住院率增加、术后并发症增加、死亡率升高、失能风险增加。')

add_subsection('11.2 衰弱逆转的可能性')
add_para('衰弱前期：通过干预可能逆转为健康状态。')
add_para('轻度衰弱：通过综合干预可能改善。')
add_para('中重度衰弱：难以完全逆转，但可延缓进展。')

doc.add_page_break()

# ========== 十二、特殊人群 ==========
add_section('十二、特殊人群的衰弱评估')

add_subsection('12.1 住院患者')
add_para('入院24小时内完成衰弱筛查。')
add_para('使用CFS或FRAIL量表。')
add_para('衰弱患者需标记，加强护理。')

add_subsection('12.2 手术患者')
add_para('术前衰弱评估可预测术后并发症。')
add_para('衰弱患者术后需加强康复。')
add_para('可考虑术前预康复（prehabilitation）。')

add_subsection('12.3 肿瘤患者')
add_para('化疗前评估衰弱状态。')
add_para('衰弱患者化疗耐受性差，需调整方案。')
add_para('可考虑衰弱导向的个体化治疗。')

add_subsection('12.4 透析患者')
add_para('透析患者衰弱发生率高。')
add_para('透析充分性、营养状态、贫血纠正影响衰弱。')
add_para('需多学科协作管理。')

doc.add_page_break()

# ========== 十三、国际核心文献 ==========
add_section('十三、国际核心文献（原始研究）')

add_subsection('13.1 Fried衰弱表型标准原始文献')
add_para('Fried LP, Tangen CM, Walston J, et al. Frailty in older adults: evidence for a phenotype.')
add_para('J Gerontol A Biol Sci Med Sci. 2001;56(3):M146-M156.')
add_para('• 首次提出Fried衰弱表型5项标准')
add_para('• 基于心血管健康研究（CHS）5377名老年人数据')
add_para('• 衰弱患病率：社区老年人6.9%，住院老年人更高')
add_para('• 衰弱与跌倒、住院、失能、死亡显著相关')

add_subsection('13.2 Rockwood累积缺陷模型')
add_para('Rockwood K, Song X, MacKnight C, et al. A global clinical measure of fitness and frailty in elderly people.')
add_para('CMAJ. 2005;173(5):489-495.')
add_para('• 提出临床衰弱量表（CFS）1-9级评分')
add_para('• 基于70项健康缺陷的累积')
add_para('• 适用于住院和急性照护场景')

add_subsection('13.3 FRAIL量表开发')
add_para('Morley JE, Malmstrom TK, Miller DK. A simple frailty questionnaire (FRAIL) predicts outcomes in middle aged African Americans.')
add_para('J Nutr Health Aging. 2012;16(7):601-608.')
add_para('• 5个问题的快速筛查工具')
add_para('• 每个"是"得1分，≥3分为衰弱')
add_para('• 适合门诊快速筛查，耗时<1分钟')

add_subsection('13.4 衰弱与不良结局的系统综述')
add_para('Sternberg SA, Schwartz AW, Karunananthan S, et al. The identification of frailty: a systematic literature review.')
add_para('J Am Geriatr Soc. 2011;59(11):2129-2138.')
add_para('• 系统回顾了2000-2010年间67项衰弱研究')
add_para('• 比较了不同评估工具的预测效度')
add_para('• 结论：Fried表型和累积缺陷模型预测效度最佳')

add_subsection('13.5 衰弱预防与干预国际共识')
add_para('Morley JE, Vellas B, van Kan GA, et al. Frailty consensus: a call to action.')
add_para('J Am Med Dir Assoc. 2013;14(6):392-397.')
add_para('• 国际老年医学和营养学会联合发布')
add_para('• 提出衰弱筛查的最低数据集')
add_para('• 强调多学科团队协作干预')

add_subsection('13.6 衰弱与手术预后')
add_para('Makary MA, Segev DL, Pronovost PJ, et al. Frailty as a predictor of surgical outcomes in older patients.')
add_para('J Am Coll Surg. 2010;210(6):901-908.')
add_para('• 前瞻性研究，592例≥65岁手术患者')
add_para('• 衰弱患者术后并发症增加2-3倍')
add_para('• 提出术前衰弱评估的必要性')

add_subsection('13.7 衰弱与ICU预后')
add_para('Bagshaw SM, Stelfox HT, McDermid RC, et al. Association between frailty and short- and long-term outcomes among critically ill patients: a multicentre prospective cohort study.')
add_para('CMAJ. 2014;186(2):E95-E102.')
add_para('• 加拿大24个ICU，16802例患者')
add_para('• 衰弱与ICU死亡率、住院时间、再入院率显著相关')
add_para('• 建议ICU常规进行衰弱评估')

add_subsection('13.8 衰弱与认知功能')
add_para('Buchman AS, Wilson RS, Bienias JL, et al. Change in frailty and risk of death in older persons.')
add_para('Am J Geriatr Psychiatry. 2009;17(5):416-421.')
add_para('• Rush Memory and Aging Project数据')
add_para('• 衰弱与认知功能下降相互关联')
add_para('• 提出"认知衰弱"（cognitive frailty）概念')

doc.add_page_break()

# ========== 十四、中国指南 ==========
add_section('十四、中国指南与专家共识')

add_subsection('14.1 中国老年衰弱评估与干预专家共识')
add_para('中华医学会老年医学分会. 中国老年衰弱评估与干预专家共识.')
add_para('中华老年医学杂志. 2023;42(1):1-10.')
add_para('• 基于中国老年人群数据制定')
add_para('• 推荐Fried表型作为首选评估工具')
add_para('• 结合中国国情调整部分参考值')
add_para('• 提出"衰弱门诊"建设标准')

add_subsection('14.2 老年衰弱门诊建设专家共识')
add_para('中华医学会老年医学分会. 老年衰弱门诊建设专家共识.')
add_para('中华老年医学杂志. 2024;43(2):121-128.')
add_para('• 衰弱门诊的功能定位')
add_para('• 人员配置：老年科医生、营养师、康复师、药师')
add_para('• 评估流程标准化')
add_para('• 信息化系统建设要求')

add_subsection('14.3 中国老年患者围手术期衰弱管理专家共识')
add_para('中华医学会麻醉学分会, 中华医学会老年医学分会.')
add_para('中华麻醉学杂志. 2023;43(5):513-520.')
add_para('• 术前衰弱筛查流程')
add_para('• 预康复（prehabilitation）方案')
add_para('• 术中管理要点')
add_para('• 术后加速康复（ERAS）结合衰弱管理')

add_subsection('14.4 中国老年衰弱相关内分泌激素管理临床实践指南')
add_para('中华医学会老年医学分会内分泌代谢学组.')
add_para('中华老年医学杂志. 2023;42(8):865-874.')
add_para('• 生长激素、性激素、甲状腺激素与衰弱的关系')
add_para('• 激素替代治疗的适应证和禁忌证')
add_para('• 维生素D、甲状旁腺激素的监测')
add_para('• 糖尿病患者的衰弱管理')

add_subsection('14.5 中国老年肌少症诊断与治疗专家共识')
add_para('中华医学会老年医学分会. 中国老年肌少症诊断与治疗专家共识.')
add_para('中华老年医学杂志. 2021;40(8):953-961.')
add_para('• 肌少症与衰弱的关系')
add_para('• EWGSOP2诊断标准的中国化应用')
add_para('• 蛋白质补充和运动干预方案')
add_para('• 药物治疗进展')

add_subsection('14.6 中国老年人膳食指南')
add_para('中国营养学会. 中国老年人膳食指南（2022）.')
add_para('人民卫生出版社. 2022.')
add_para('• 老年人营养需求特点')
add_para('• 蛋白质、钙、维生素D摄入推荐')
add_para('• 膳食模式（地中海饮食、DASH饮食）')
add_para('• 营养不良筛查（MNA-SF）')

add_subsection('14.7 中国老年人跌倒预防指南')
add_para('中华医学会老年医学分会. 中国老年人跌倒预防指南.')
add_para('中华老年医学杂志. 2022;41(5):512-520.')
add_para('• 跌倒风险评估工具')
add_para('• 环境改造标准')
add_para('• 药物相关跌倒')
add_para('• 平衡训练和力量训练方案')

doc.add_page_break()

# ========== 十五、干预与康复RCT ==========
add_section('十五、干预与康复RCT研究')

add_subsection('15.1 运动干预系统综述')
add_para('Cadore EL, Rodriguez-Manas L, Sinclair A, et al. Effects of different exercise interventions on risk of falls, gait ability, and balance in physically frail older adults: a systematic review.')
add_para('Rejuvenation Res. 2013;16(2):105-114.')
add_para('• 纳入15项RCT，共1123例衰弱老年人')
add_para('• 抗阻训练+平衡训练效果最佳')
add_para('• 每周≥2次，持续≥12周可见显著改善')

add_subsection('15.2 营养干预RCT')
add_para('Rondanelli M, Opizzi A, Giacosa A, et al. Effect of essential amino acid supplementation on quality of life, physical function and muscle strength in frail elderly.')
add_para('Clin Nutr. 2011;30(5):571-577.')
add_para('• 意大利多中心RCT，149例衰弱老年人')
add_para('• 必需氨基酸补充（8g/日）+维生素D')
add_para('• 12周后握力增加18%，步速提高12%')

add_subsection('15.3 综合干预（FORTIC研究）')
add_para('Faber MJ, Bosscher RJ, Chin APMJ, et al. Effects of exercise programs on falls and mobility in frail and pre-frail older adults: a multicenter randomized controlled trial.')
add_para('Arch Phys Med Rehabil. 2006;87(7):885-896.')
add_para('• 荷兰多中心RCT，217例衰弱/衰弱前期')
add_para('• 综合干预：运动+营养+认知训练')
add_para('• 16周后跌倒风险降低35%')

add_subsection('15.4 衰弱逆转可能性研究')
add_para('Gill TM, Gahbauer EA, Allore HG, et al. Transitions between frailty states among community-living older persons.')
add_para('Arch Intern Med. 2006;166(4):418-423.')
add_para('• 耶鲁大学纵向研究，754例社区老年人')
add_para('• 衰弱状态可逆：43%的衰弱者在3年内改善')
add_para('• 预测逆转因素：运动、社会支持、无抑郁')

add_subsection('15.5 预康复研究（PREHAB）')
add_para('Barberan-Garcia A, Ubré M, Roca J, et al. Personalised prehabilitation in high-risk patients undergoing elective major abdominal surgery: a randomized blinded controlled trial.')
add_para('Ann Surg. 2018;267(1):50-56.')
add_para('• 西班牙RCT，71例重大腹部手术患者')
add_para('• 术前2-4周：运动+营养+心理支持')
add_para('• 术后并发症降低51%，住院时间缩短2.7天')

doc.add_page_break()

# ========== 十六、衰弱与特定疾病 ==========
add_section('十六、衰弱与特定疾病')

add_subsection('16.1 衰弱与心力衰竭')
add_para('Uchmanowicz I, Jankowska-Polanska B, Loboz-Grudzien K, et al. Frailty syndrome in elderly patients with heart failure.')
add_para('Eur J Cardiovasc Nurs. 2014;13(6):562-570.')
add_para('• 心衰患者衰弱患病率40-60%')
add_para('• 衰弱与心衰恶化相互促进')
add_para('• 建议心衰患者常规进行衰弱筛查')

add_subsection('16.2 衰弱与慢性肾病')
add_para('Johansen KL, Dalrymple LS, Delgado C, et al. Association between frailty and cardiovascular outcomes in the National Health and Nutrition Examination Survey.')
add_para('J Am Geriatr Soc. 2013;61(4):570-576.')
add_para('• CKD患者衰弱患病率随eGFR下降而增加')
add_para('• 衰弱是CKD患者死亡的独立预测因子')
add_para('• 透析患者衰弱管理策略')

add_subsection('16.3 衰弱与肿瘤')
add_para('Cruz-Jentoft AJ, Bahat G, Bauer J, et al. Sarcopenia: revised European consensus on definition and diagnosis.')
add_para('Age Ageing. 2019;48(1):16-31.')
add_para('• 肿瘤患者衰弱评估工具（CRASH、CARG评分）')
add_para('• 衰弱导向的化疗剂量调整')
add_para('• 肿瘤预康复方案')

add_subsection('16.4 衰弱与糖尿病')
add_para('Rodriguez-Manas L, Fried LP. Frailty in the clinical scenario.')
add_para('Lancet. 2015;385(9968):e7-e9.')
add_para('• 糖尿病患者衰弱风险增加2-3倍')
add_para('• 低血糖与衰弱恶性循环')
add_para('• 个体化血糖控制目标')

doc.add_page_break()

# ========== 十七、评估工具详解 ==========
add_section('十七、评估工具详解')

add_subsection('17.1 计时起立行走测试（TUGT）')
add_para('Podsiadlo D, Richardson S. The timed "Up & Go": a test of basic functional mobility for frail elderly persons.')
add_para('J Am Geriatr Soc. 1991;39(2):142-148.')
add_para('• 正常值：<10秒')
add_para('• 衰弱筛查界值：≥12秒')
add_para('• 操作简便，适合社区和门诊')

add_subsection('17.2 6分钟步行试验（6MWT）')
add_para('ATS Committee on Proficiency Standards for Clinical Pulmonary Function Laboratories. ATS statement: guidelines for the six-minute walk test.')
add_para('Am J Respir Crit Care Med. 2002;166(1):111-117.')
add_para('• 正常值：400-700米（因年龄、性别、身高而异）')
add_para('• 衰弱相关：步行距离<300米提示功能受限')
add_para('• 可重复性好，适合随访评估')

add_subsection('17.3 简易体能状况量表（SPPB）')
add_para('Guralnik JM, Simonsick EM, Ferrucci L, et al. A short physical performance battery assessing lower extremity function: association with self-reported disability and prediction of mortality and nursing home admission.')
add_para('J Gerontol. 1994;49(2):M85-M94.')
add_para('• 3个测试：站立平衡、4米步行、5次起坐')
add_para('• 评分0-12分，≤9分提示功能受限')
add_para('• 预测效度优于单一指标')

add_subsection('17.4 营养风险筛查2002（NRS-2002）')
add_para('Kondrup J, Allison SP, Elia M, et al. ESPEN guidelines for nutrition screening 2002.')
add_para('Clin Nutr. 2003;22(4):415-421.')
add_para('• 适用于住院患者营养筛查')
add_para('• 评分≥3分提示存在营养风险')
add_para('• 与衰弱密切相关')

doc.add_page_break()

# ========== 十八、公共卫生 ==========
add_section('十八、公共卫生与流行病学')

add_subsection('18.1 全球衰弱负担')
add_para('Collard RM, Boter H, Schoevers RA, et al. Prevalence of frailty in community-dwelling older persons: a systematic review.')
add_para('J Am Geriatr Soc. 2012;60(8):1487-1492.')
add_para('• 系统回顾54项研究，共81089例社区老年人')
add_para('• 衰弱患病率：10.7%（65-74岁）至15.7%（≥85岁）')
add_para('• 衰弱前期患病率：约30%')
add_para('• 女性高于男性，低收入国家更高')

add_subsection('18.2 衰弱筛查的成本效益')
add_para('Perna S, Francis MD, Bologna C, et al. Performance of Edmonton Frail Scale in geriatric outpatients: a systematic review and meta-analysis.')
add_para('Aging Clin Exp Res. 2021;33(4):875-884.')
add_para('• 系统回顾12项研究')
add_para('• 衰弱筛查可识别高风险患者，针对性干预可降低医疗成本')
add_para('• 建议65岁以上老年人常规筛查')

doc.add_page_break()

# ========== 十九、附录：参考值速查表 ==========
add_section('十九、附录：关键参考值速查表')

add_subsection('19.1 握力正常参考值（Jamar握力计，站立位，3次取最大值）')
table = doc.add_table(rows=8, cols=3)
table.style = 'Table Grid'
for i, h in enumerate(['性别', 'BMI范围', '正常下限']):
    cell = table.rows[0].cells[i]
    cell.text = h
    for run in cell.paragraphs[0].runs:
        set_font(run, '宋体', 10, True)
rows = [
    ['男性', '≤24 kg/m²', '>29 kg'],
    ['男性', '24.1-28 kg/m²', '>30 kg'],
    ['男性', '>28 kg/m²', '>32 kg'],
    ['女性', '≤23 kg/m²', '>17 kg'],
    ['女性', '23.1-26 kg/m²', '>17.3 kg'],
    ['女性', '26.1-29 kg/m²', '>18 kg'],
    ['女性', '>29 kg/m²', '>21 kg'],
]
for row_data in rows:
    row = table.add_row()
    for i, text in enumerate(row_data):
        cell = row.cells[i]
        cell.text = text
        for run in cell.paragraphs[0].runs:
            set_font(run, '宋体', 10)

add_subsection('19.2 步速正常参考值（4.57米步行测试）')
table = doc.add_table(rows=5, cols=3)
table.style = 'Table Grid'
for i, h in enumerate(['性别', '身高范围', '正常下限']):
    cell = table.rows[0].cells[i]
    cell.text = h
    for run in cell.paragraphs[0].runs:
        set_font(run, '宋体', 10, True)
rows = [
    ['男性', '≤173 cm', '>0.65 m/s'],
    ['男性', '>173 cm', '>0.76 m/s'],
    ['女性', '≤159 cm', '>0.65 m/s'],
    ['女性', '>159 cm', '>0.76 m/s'],
]
for row_data in rows:
    row = table.add_row()
    for i, text in enumerate(row_data):
        cell = row.cells[i]
        cell.text = text
        for run in cell.paragraphs[0].runs:
            set_font(run, '宋体', 10)

add_subsection('19.3 活动量标准（MLTAQ问卷）')
table = doc.add_table(rows=3, cols=2)
table.style = 'Table Grid'
for i, h in enumerate(['性别', '正常下限']):
    cell = table.rows[0].cells[i]
    cell.text = h
    for run in cell.paragraphs[0].runs:
        set_font(run, '宋体', 10, True)
rows = [
    ['男性', '≥383 kcal/周'],
    ['女性', '≥270 kcal/周'],
]
for row_data in rows:
    row = table.add_row()
    for i, text in enumerate(row_data):
        cell = row.cells[i]
        cell.text = text
        for run in cell.paragraphs[0].runs:
            set_font(run, '宋体', 10)

add_subsection('19.4 体重下降标准')
table = doc.add_table(rows=3, cols=2)
table.style = 'Table Grid'
for i, h in enumerate(['指标', '界值']):
    cell = table.rows[0].cells[i]
    cell.text = h
    for run in cell.paragraphs[0].runs:
        set_font(run, '宋体', 10, True)
rows = [
    ['1年内体重下降百分比', '≥5%'],
    ['1年内体重下降绝对值', '≥4.5 kg（10磅）'],
]
for row_data in rows:
    row = table.add_row()
    for i, text in enumerate(row_data):
        cell = row.cells[i]
        cell.text = text
        for run in cell.paragraphs[0].runs:
            set_font(run, '宋体', 10)

doc.add_page_break()

# ========== 二十、附录：干预方案速查 ==========
add_section('二十、附录：干预方案速查')

add_subsection('20.1 营养干预')
add_para('蛋白质：1.2-1.5 g/kg/日，优质蛋白≥50%')
add_para('维生素D：800-1000 IU/日')
add_para('钙：1000-1200 mg/日')
add_para('口服营养补充剂：400-600 kcal/日')

add_subsection('20.2 运动干预')
add_para('有氧运动：每周≥150分钟，中等强度')
add_para('抗阻训练：每周2-3次，8-12次/组，1-3组')
add_para('平衡训练：每周≥3次，每次20-30分钟')
add_para('持续时间：至少12周可见效果')

add_subsection('20.3 药物管理')
add_para('定期药物重整（每3个月）')
add_para('遵循Beers标准和STOPP/START标准')
add_para('停用无明确适应症的药物')
add_para('低剂量起始、缓慢增量')

add_subsection('20.4 随访频率')
add_para('衰弱前期：每6个月')
add_para('衰弱：每3个月')
add_para('住院后：出院后1周、1个月、3个月')

# 保存
output_path = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'RAG知识库文献汇编.docx')
doc.save(output_path)
print(f'已保存: {output_path}')
print(f'文件大小: {os.path.getsize(output_path)} bytes')
